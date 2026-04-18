from dotenv import load_dotenv
import html as _html
import base64
import streamlit as st
import streamlit.components.v1 as components
from PyPDF2 import PdfReader
import pandas as pd
from docx import Document
from PIL import Image, ImageEnhance, ImageOps
import pytesseract
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.messages import HumanMessage
from langchain_experimental.agents import create_csv_agent
import re
from datetime import datetime
import json
import io
import csv
import os
import shutil
import tempfile
import time
import requests
from youtube_transcript_api import YouTubeTranscriptApi
from bs4 import BeautifulSoup
from fpdf import FPDF

try:
    import tabulate  # noqa: F401
    HAS_TABULATE = True
except ImportError:
    HAS_TABULATE = False

# Load environment variables
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")

SUPPORTED_MODEL_ALIASES = {
    "gemini-1.5-flash-latest": "gemini-2.5-flash",
    "gemini-1.5-pro-latest": "gemini-2.5-pro",
}

AVAILABLE_GEMINI_MODELS = [
    "gemini-2.5-flash",
    "gemini-2.5-pro",
    "gemini-2.0-flash",
    "gemini-1.5-flash",
    "gemini-1.5-pro",
]

MODEL_TIMEOUT_SECONDS = 25
MODEL_MAX_RETRIES = 1


def normalize_model_name(model_name):
    return SUPPORTED_MODEL_ALIASES.get(model_name, model_name)


def _get_model_fallback_order(primary_model, fallback_models):
    """Return de-duplicated model order: primary first, then fallbacks."""
    ordered = [normalize_model_name(primary_model)]
    for model_name in fallback_models or []:
        normalized = normalize_model_name(model_name)
        if normalized not in ordered:
            ordered.append(normalized)
    return ordered

@st.cache_resource(show_spinner="⏳ Loading AI model (first run only)...")
def get_embeddings():
    """Cache the embedding model so it loads only once per server run."""
    return HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

def generate_chat_pdf(chat_history, file_name=None, summary_text=None):
    """Generate a beautifully formatted PDF of the chat history."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # ---- Header banner ----
    pdf.set_fill_color(102, 126, 234)
    pdf.rect(0, 0, 210, 38, 'F')
    pdf.set_text_color(255, 255, 255)
    pdf.set_font('Helvetica', 'B', 22)
    pdf.set_xy(0, 8)
    pdf.cell(210, 10, 'SmartAI ChatBot', align='C', ln=True)
    pdf.set_font('Helvetica', '', 11)
    pdf.set_xy(0, 20)
    pdf.cell(210, 8, 'Chat History Report', align='C', ln=True)

    # ---- Meta info ----
    pdf.set_xy(0, 30)
    pdf.set_font('Helvetica', '', 8)
    pdf.cell(210, 6, f'Generated: {datetime.now().strftime("%B %d, %Y  %H:%M")}', align='C', ln=True)
    pdf.ln(6)

    if file_name:
        pdf.set_text_color(80, 80, 80)
        pdf.set_font('Helvetica', 'I', 9)
        pdf.set_x(15)
        source_label = file_name if len(file_name) <= 80 else file_name[:77] + '...'
        pdf.cell(0, 6, f'Source: {source_label}', ln=True)
        pdf.ln(2)

    pdf.set_draw_color(200, 200, 200)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(4)

    # Optional summary block (used for URL/Video exports)
    if summary_text:
        pdf.set_text_color(6, 95, 70)
        pdf.set_font('Helvetica', 'B', 10)
        pdf.set_x(15)
        pdf.cell(0, 7, 'AI Summary:', ln=True)
        pdf.set_font('Helvetica', '', 10)
        pdf.set_fill_color(236, 253, 245)
        pdf.set_text_color(30, 41, 59)
        s_text = summary_text.encode('latin-1', errors='replace').decode('latin-1')
        pdf.set_x(20)
        pdf.multi_cell(170, 6, s_text, fill=True)
        pdf.ln(4)

        pdf.set_draw_color(220, 220, 220)
        pdf.line(15, pdf.get_y(), 195, pdf.get_y())
        pdf.ln(4)

    # ---- Q&A entries ----
    for i, entry in enumerate(chat_history, 1):
        # Entry number + timestamp row
        pdf.set_fill_color(102, 126, 234)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font('Helvetica', 'B', 9)
        pdf.set_x(15)
        pdf.cell(18, 6, f'  #{i}', fill=True, ln=False)

        ts = entry.get('timestamp', '')
        try:
            ts = datetime.fromisoformat(ts).strftime('%d %b %Y, %H:%M')
        except Exception:
            pass
        pdf.set_fill_color(240, 240, 240)
        pdf.set_text_color(100, 100, 100)
        pdf.set_font('Helvetica', 'I', 8)
        pdf.cell(0, 6, f'  {ts}  |  {entry.get("file_type", "").upper()}', fill=True, ln=True)
        pdf.ln(2)

        # Question box
        pdf.set_text_color(40, 50, 120)
        pdf.set_font('Helvetica', 'B', 9)
        pdf.set_x(15)
        pdf.cell(0, 6, 'You asked:', ln=True)
        pdf.set_font('Helvetica', '', 10)
        pdf.set_fill_color(242, 245, 255)
        pdf.set_text_color(30, 41, 59)
        q_text = entry.get('question', '').encode('latin-1', errors='replace').decode('latin-1')
        pdf.set_x(20)
        pdf.multi_cell(170, 6, q_text, fill=True)
        pdf.ln(2)

        # Answer box
        pdf.set_text_color(6, 95, 70)
        pdf.set_font('Helvetica', 'B', 9)
        pdf.set_x(15)
        pdf.cell(0, 6, 'SmartAI Answer:', ln=True)
        pdf.set_font('Helvetica', '', 10)
        pdf.set_fill_color(240, 253, 244)
        pdf.set_text_color(30, 41, 59)
        a_text = entry.get('answer', '').encode('latin-1', errors='replace').decode('latin-1')
        pdf.set_x(20)
        pdf.multi_cell(170, 6, a_text, fill=True)
        pdf.ln(4)

        # Separator line
        pdf.set_draw_color(220, 220, 220)
        pdf.line(15, pdf.get_y(), 195, pdf.get_y())
        pdf.ln(4)

    # ---- Footer ----
    pdf.set_y(-18)
    pdf.set_draw_color(102, 126, 234)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.set_font('Helvetica', 'I', 8)
    pdf.set_text_color(150, 150, 150)
    pdf.cell(0, 8, 'SmartAI ChatBot  |  Powered by Google Gemini & LangChain', align='C')

    return bytes(pdf.output())

# Page configuration
st.set_page_config(
    page_title="SmartAI ChatBot - Multi-Format AI Assistant",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)


def _get_custom_hero_image_data_uri():
    """Return data URI for a user-provided hero image, if present."""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    candidates = [
        "AI-Face.png",
        os.path.join("assets", "AI-Face.png"),
    ]
    mime_map = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
    }

    for rel_path in candidates:
        full_path = os.path.join(base_dir, rel_path)
        if not os.path.exists(full_path):
            continue
        ext = os.path.splitext(full_path)[1].lower()
        mime = mime_map.get(ext)
        if not mime:
            continue
        with open(full_path, "rb") as f:
            encoded = base64.b64encode(f.read()).decode("ascii")
        return f"data:{mime};base64,{encoded}"

    return None

# Custom CSS for dark, production-style UI
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@400;500;600;700;800&family=Sora:wght@500;600;700;800&display=swap');
    @import url('https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.5.2/css/all.min.css');

    :root {
        --bg-0: #03051a;
        --bg-1: #070a22;
        --bg-2: #111532;
        --panel: rgba(11, 16, 32, 0.72);
        --panel-strong: rgba(18, 25, 45, 0.92);
        --panel-soft: rgba(255, 255, 255, 0.06);
        --border: rgba(190, 174, 255, 0.14);
        --text-0: #f8fafc;
        --text-1: #e2e8f0;
        --text-2: #b9c5d8;
        --brand-1: #8b5cf6;
        --brand-2: #d0b3ff;
        --brand-3: #7dd3fc;
        --danger: #ef4444;
        --shadow-xl: 0 28px 60px rgba(2, 6, 23, 0.48);
        --shadow-md: 0 12px 32px rgba(2, 6, 23, 0.32);
        --radius-xl: 28px;
        --radius-lg: 20px;
        --radius-md: 14px;
    }

        html[data-theme="light"] {
            --bg-0: #f4f7ff;
            --bg-1: #eef2ff;
            --bg-2: #e4eafc;
            --panel: rgba(255, 255, 255, 0.86);
            --panel-strong: rgba(255, 255, 255, 0.96);
            --panel-soft: rgba(99, 102, 241, 0.06);
            --border: rgba(99, 102, 241, 0.18);
            --text-0: #0f172a;
            --text-1: #334155;
            --text-2: #64748b;
            --shadow-xl: 0 24px 52px rgba(15, 23, 42, 0.12);
            --shadow-md: 0 10px 24px rgba(15, 23, 42, 0.08);
        }

    html, body, [class*="css"] {
        font-family: 'Outfit', 'Segoe UI', Tahoma, sans-serif;
        font-size: 16.5px;
        -webkit-font-smoothing: antialiased;
        text-rendering: optimizeLegibility;
    }

    .stApp {
        color: var(--text-0);
        background:
            radial-gradient(circle at 16% 10%, rgba(196, 181, 253, 0.18), transparent 18%),
            radial-gradient(circle at 78% 14%, rgba(139, 92, 246, 0.24), transparent 22%),
            radial-gradient(circle at 50% 34%, rgba(67, 56, 202, 0.16), transparent 28%),
            radial-gradient(circle at 52% 112%, rgba(125, 211, 252, 0.09), transparent 26%),
            linear-gradient(180deg, #030412 0%, #06081d 22%, #090d28 52%, #05081c 76%, #02030c 100%);
        min-height: 100vh;
    }

    html[data-theme="light"] .stApp {
        background:
            radial-gradient(circle at 16% 10%, rgba(129, 140, 248, 0.12), transparent 18%),
            radial-gradient(circle at 78% 14%, rgba(168, 85, 247, 0.1), transparent 22%),
            radial-gradient(circle at 50% 34%, rgba(59, 130, 246, 0.08), transparent 28%),
            radial-gradient(circle at 52% 112%, rgba(56, 189, 248, 0.06), transparent 26%),
            linear-gradient(180deg, #f9fbff 0%, #f3f6ff 22%, #edf2ff 52%, #eef4ff 76%, #f8fbff 100%);
    }

    .stApp::before {
        content: "";
        position: fixed;
        inset: 0;
        background:
            radial-gradient(circle at 50% -10%, rgba(216, 180, 254, 0.12), transparent 34%),
            linear-gradient(180deg, rgba(255,255,255,0.028) 0, rgba(255,255,255,0) 140px),
            linear-gradient(90deg, rgba(192, 170, 255, 0.035) 1px, transparent 1px),
            linear-gradient(rgba(192, 170, 255, 0.03) 1px, transparent 1px);
        background-size: auto, auto, 96px 96px, 96px 96px;
        pointer-events: none;
        mask-image: radial-gradient(circle at center, black 58%, transparent 100%);
        opacity: 0.9;
    }

    html[data-theme="light"] .stApp::before {
        background:
            radial-gradient(circle at 50% -10%, rgba(167, 139, 250, 0.08), transparent 34%),
            linear-gradient(180deg, rgba(99, 102, 241, 0.04) 0, rgba(255,255,255,0) 140px),
            linear-gradient(90deg, rgba(99, 102, 241, 0.045) 1px, transparent 1px),
            linear-gradient(rgba(99, 102, 241, 0.04) 1px, transparent 1px);
        opacity: 0.7;
    }

    .stApp::after {
        content: "";
        position: fixed;
        inset: 0;
        pointer-events: none;
        background: radial-gradient(circle at 50% 0%, rgba(255,255,255,0.04), transparent 42%);
        mix-blend-mode: screen;
        opacity: 0.35;
    }

    html[data-theme="light"] .stApp::after {
        background: radial-gradient(circle at 50% 0%, rgba(255,255,255,0.48), transparent 42%);
        mix-blend-mode: normal;
        opacity: 0.28;
    }

    [data-testid="stAppViewContainer"] {
        background: transparent;
    }

    [data-testid="stHeader"] {
        background: transparent;
    }

    [data-testid="stSidebar"] {
        background:
            linear-gradient(180deg, rgba(8, 13, 27, 0.95) 0%, rgba(9, 15, 29, 0.92) 100%);
        border-right: 1px solid var(--border);
        box-shadow: inset -1px 0 0 rgba(255, 255, 255, 0.04);
    }

    html[data-theme="light"] [data-testid="stSidebar"] {
        background: linear-gradient(180deg, rgba(246, 248, 255, 0.96) 0%, rgba(238, 242, 255, 0.94) 100%);
        box-shadow: inset -1px 0 0 rgba(99, 102, 241, 0.06);
    }

    [data-testid="stSidebar"] > div {
        background: transparent;
    }

    [data-testid="stSidebar"] * {
        color: var(--text-0);
    }

    .sidebar-hero,
    .sidebar-card,
    .sidebar-meta-card {
        position: relative;
        overflow: hidden;
        transition: transform 0.22s ease, box-shadow 0.22s ease, border-color 0.22s ease, background 0.22s ease;
    }

    .sidebar-hero::before,
    .sidebar-card::before,
    .sidebar-meta-card::before {
        content: "";
        position: absolute;
        inset: 0;
        background: linear-gradient(120deg, rgba(255, 255, 255, 0.08), transparent 35%, transparent 65%, rgba(6, 182, 212, 0.08));
        opacity: 0;
        transition: opacity 0.22s ease;
        pointer-events: none;
    }

    .sidebar-hero:hover,
    .sidebar-card:hover,
    .sidebar-meta-card:hover {
        transform: translateY(-2px);
        border-color: rgba(125, 211, 252, 0.28) !important;
        box-shadow: 0 18px 36px rgba(2, 6, 23, 0.38) !important;
    }

    .sidebar-hero:hover::before,
    .sidebar-card:hover::before,
    .sidebar-meta-card:hover::before {
        opacity: 1;
    }

    .sidebar-hero h2,
    .sidebar-card h4,
    .sidebar-meta-card strong {
        position: relative;
        z-index: 1;
    }

    .sidebar-hero p,
    .sidebar-card p,
    .sidebar-meta-card small {
        position: relative;
        z-index: 1;
    }

    /* ── Golden Meta Card ─────────────────────────────── */
    .sidebar-meta-card {
        background:
            url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='200' height='200' viewBox='0 0 200 200'%3E%3Cg stroke='%23fde68a' stroke-width='0.6' opacity='0.18' fill='none'%3E%3Cline x1='30' y1='50' x2='100' y2='35'/%3E%3Cline x1='30' y1='50' x2='100' y2='80'/%3E%3Cline x1='30' y1='50' x2='100' y2='125'/%3E%3Cline x1='30' y1='50' x2='100' y2='165'/%3E%3Cline x1='30' y1='100' x2='100' y2='35'/%3E%3Cline x1='30' y1='100' x2='100' y2='80'/%3E%3Cline x1='30' y1='100' x2='100' y2='125'/%3E%3Cline x1='30' y1='100' x2='100' y2='165'/%3E%3Cline x1='30' y1='155' x2='100' y2='35'/%3E%3Cline x1='30' y1='155' x2='100' y2='80'/%3E%3Cline x1='30' y1='155' x2='100' y2='125'/%3E%3Cline x1='30' y1='155' x2='100' y2='165'/%3E%3C/g%3E%3Cg stroke='%23fbbf24' stroke-width='0.6' opacity='0.18' fill='none'%3E%3Cline x1='100' y1='35' x2='170' y2='75'/%3E%3Cline x1='100' y1='35' x2='170' y2='130'/%3E%3Cline x1='100' y1='80' x2='170' y2='75'/%3E%3Cline x1='100' y1='80' x2='170' y2='130'/%3E%3Cline x1='100' y1='125' x2='170' y2='75'/%3E%3Cline x1='100' y1='125' x2='170' y2='130'/%3E%3Cline x1='100' y1='165' x2='170' y2='75'/%3E%3Cline x1='100' y1='165' x2='170' y2='130'/%3E%3C/g%3E%3Ccircle cx='30' cy='50' r='2' fill='%23fde68a' opacity='0.38'/%3E%3Ccircle cx='30' cy='100' r='2' fill='%23fde68a' opacity='0.38'/%3E%3Ccircle cx='30' cy='155' r='2' fill='%23fde68a' opacity='0.38'/%3E%3Ccircle cx='100' cy='35' r='2' fill='%23fbbf24' opacity='0.35'/%3E%3Ccircle cx='100' cy='80' r='2' fill='%23fbbf24' opacity='0.35'/%3E%3Ccircle cx='100' cy='125' r='2' fill='%23fbbf24' opacity='0.35'/%3E%3Ccircle cx='100' cy='165' r='2' fill='%23fbbf24' opacity='0.35'/%3E%3Ccircle cx='170' cy='75' r='2.5' fill='%23f59e0b' opacity='0.4'/%3E%3Ccircle cx='170' cy='130' r='2.5' fill='%23f59e0b' opacity='0.4'/%3E%3C/svg%3E") center/cover no-repeat,
            linear-gradient(135deg,
                rgba(180,130,20,0.18) 0%,
                rgba(140,100,10,0.22) 40%,
                rgba(80,55,5,0.26) 100%) !important;
        border: 1px solid rgba(253,230,138,0.38) !important;
        box-shadow:
            0 0 18px rgba(253,230,138,0.16),
            0 6px 28px rgba(0,0,0,0.32),
            inset 0 1px 0 rgba(255,248,180,0.14) !important;
        animation: metaCardGlow 4s ease-in-out infinite;
    }
    .sidebar-meta-card::before {
        background: linear-gradient(120deg,
            rgba(255,248,180,0.1),
            transparent 40%,
            rgba(253,230,138,0.08)) !important;
        opacity: 1 !important;
    }
    .meta-title {
        display: inline-block;
        font-size: 15px;
        font-weight: 700;
        letter-spacing: 0.02em;
        background: linear-gradient(90deg,
            #fff7c0 0%, #fde68a 18%, #fbbf24 32%,
            #ffffff 50%, #fbbf24 68%, #fde68a 82%, #fff7c0 100%);
        background-size: 250% auto;
        -webkit-background-clip: text;
        background-clip: text;
        -webkit-text-fill-color: transparent;
        animation: goldShimmer 2.6s linear infinite;
    }
    .meta-sub {
        display: block;
        font-size: 12px;
        color: #fde68a;
        margin-top: 3px;
        animation: metaSubDrift 7s ease-in-out infinite;
    }
    @keyframes goldShimmer {
        0%   { background-position: 0%   center; }
        100% { background-position: 250% center; }
    }
    @keyframes metaTitleDrift {
        0%   { transform: translate(0px,   0px);  }
        15%  { transform: translate(6px,  -8px);  }
        30%  { transform: translate(-4px, -5px);  }
        48%  { transform: translate(7px,   4px);  }
        63%  { transform: translate(-6px,  7px);  }
        78%  { transform: translate(3px,  -9px);  }
        100% { transform: translate(0px,   0px);  }
    }
    @keyframes metaSubDrift {
        0%   { transform: translate(0px,  0px);  opacity: 0.75; }
        18%  { transform: translate(-7px, 6px);  opacity: 1;    }
        35%  { transform: translate(5px,  8px);  opacity: 0.8;  }
        52%  { transform: translate(-4px,-6px);  opacity: 1;    }
        68%  { transform: translate(8px, -4px);  opacity: 0.8;  }
        84%  { transform: translate(-5px, 5px);  opacity: 1;    }
        100% { transform: translate(0px,  0px);  opacity: 0.75; }
    }
    @keyframes metaCardGlow {
        0%,100% {
            box-shadow:
                0 0 18px rgba(251,191,36,0.22),
                0 6px 28px rgba(0,0,0,0.5),
                inset 0 1px 0 rgba(253,230,138,0.18);
            border-color: rgba(251,191,36,0.46) !important;
        }
        50% {
            box-shadow:
                0 0 36px rgba(251,191,36,0.46),
                0 10px 36px rgba(0,0,0,0.56),
                inset 0 1px 0 rgba(253,230,138,0.26);
            border-color: rgba(251,191,36,0.72) !important;
        }
    }

    .guide-hero,
    .guide-step,
    .guide-step-shell,
    .guide-capability,
    .guide-panel,
    .guide-cta {
        position: relative;
        overflow: hidden;
        transition: transform 0.22s ease, box-shadow 0.22s ease, border-color 0.22s ease, background 0.22s ease;
    }

    .guide-hero::before,
    .guide-step::before,
    .guide-step-shell::before,
    .guide-capability::before,
    .guide-panel::before,
    .guide-cta::before {
        content: "";
        position: absolute;
        inset: 0;
        background: linear-gradient(125deg, rgba(255,255,255,0.08), transparent 34%, transparent 68%, rgba(208,179,255,0.12));
        opacity: 0;
        transition: opacity 0.22s ease;
        pointer-events: none;
    }

    .guide-hero:hover,
    .guide-step:hover,
    .guide-step-shell:hover,
    .guide-capability:hover,
    .guide-panel:hover,
    .guide-cta:hover {
        transform: translateY(-2px);
        box-shadow: 0 16px 30px rgba(2, 6, 23, 0.32);
        border-color: rgba(208, 179, 255, 0.28) !important;
    }

    .guide-hero:hover::before,
    .guide-step:hover::before,
    .guide-step-shell:hover::before,
    .guide-capability:hover::before,
    .guide-panel:hover::before,
    .guide-cta:hover::before {
        opacity: 1;
    }

    .guide-grid {
        display: flex;
        flex-direction: column;
        gap: 0.9rem;
    }

    .guide-loader {
        margin: 0.75rem 0 0.95rem 0;
        padding: 0.75rem 0.7rem;
        border-radius: 12px;
        background: rgba(255,255,255,0.03);
        border: 1px solid rgba(255,255,255,0.08);
    }

    .guide-loader-track {
        display: grid;
        grid-template-columns: repeat(3, 1fr);
        gap: 0.5rem;
        align-items: center;
    }

    .guide-loader-node {
        position: relative;
        display: flex;
        align-items: center;
        justify-content: center;
        min-height: 2.4rem;
        border-radius: 11px;
        border: 1px solid rgba(255,255,255,0.08);
        background: rgba(255,255,255,0.03);
        color: #cbd5e1;
        font-size: 0.75rem;
        font-weight: 700;
        letter-spacing: 0.02em;
    }

    .guide-loader-node.active {
        border-color: rgba(125, 211, 252, 0.44);
        background: linear-gradient(135deg, rgba(79,70,229,0.4), rgba(8,145,178,0.26));
        color: #f8fafc;
        box-shadow: 0 0 0 1px rgba(255,255,255,0.08), 0 12px 20px rgba(8, 145, 178, 0.18);
        animation: loaderPulse 1.8s ease-in-out infinite;
    }

    .guide-loader-node.done {
        border-color: rgba(74, 222, 128, 0.45);
        background: linear-gradient(135deg, rgba(22,163,74,0.25), rgba(6,182,212,0.2));
        color: #dcfce7;
    }

    .guide-capability.featured {
        border-left: 3px solid #22c55e;
        background: linear-gradient(135deg, rgba(34,197,94,0.13), rgba(6,182,212,0.08));
    }

    @keyframes loaderPulse {
        0%, 100% {
            box-shadow: 0 0 0 1px rgba(255,255,255,0.08), 0 12px 20px rgba(8, 145, 178, 0.18);
            transform: translateY(0px);
        }
        50% {
            box-shadow: 0 0 0 1px rgba(186, 230, 253, 0.32), 0 16px 26px rgba(8, 145, 178, 0.26);
            transform: translateY(-1px);
        }
    }

    .guide-step {
        padding: 0.9rem;
        border-radius: 14px;
        background: rgba(255,255,255,0.04);
        border: 1px solid rgba(255,255,255,0.08);
    }

    .guide-step-index {
        display: inline-flex;
        align-items: center;
        justify-content: center;
        width: 1.55rem;
        height: 1.55rem;
        margin-bottom: 0.55rem;
        border-radius: 999px;
        background: linear-gradient(135deg, rgba(79,70,229,0.95), rgba(6,182,212,0.85));
        color: #ffffff;
        font-size: 0.76rem;
        font-weight: 800;
    }

    .guide-step-shell {
        position: relative;
        padding: 1.25rem 1rem 1.15rem 1rem;
        border-radius: 22px;
        background:
            radial-gradient(circle at 50% -8%, rgba(255,255,255,0.06), transparent 26%),
            radial-gradient(circle at 50% 110%, rgba(160, 0, 40, 0.09), transparent 28%),
            linear-gradient(180deg, rgba(14, 6, 10, 0.99), rgba(8, 3, 6, 1));
        border: 1px solid rgba(180, 40, 70, 0.28);
        box-shadow: inset 0 1px 0 rgba(255,255,255,0.04), 0 22px 42px rgba(2, 0, 6, 0.35);
        min-height: 198px;
        text-align: center;
        overflow: hidden;
    }

    .guide-step-shell::before {
        content: "";
        position: absolute;
        inset: 0;
        border-radius: inherit;
        background:
            radial-gradient(circle, rgba(200, 30, 60, 0.55) 0.7px, transparent 0.9px);
        background-size: 5px 5px;
        opacity: 0.28;
        pointer-events: none;
        mix-blend-mode: screen;
    }

    .guide-step-shell::after {
        content: "";
        position: absolute;
        inset: 0;
        border-radius: inherit;
        background:
            radial-gradient(circle at 50% 8%, rgba(255,255,255,0.12), transparent 18%),
            radial-gradient(circle at 15% 12%, rgba(180, 30, 60, 0.18), transparent 16%),
            radial-gradient(circle at 85% 18%, rgba(180, 30, 60, 0.18), transparent 18%),
            linear-gradient(180deg, rgba(255,255,255,0.02), transparent 30%, rgba(160, 10, 30, 0.05) 100%);
        opacity: 0.75;
        pointer-events: none;
    }

    .guide-step-pill {
        position: relative;
        z-index: 1;
        display: inline-flex;
        align-items: center;
        justify-content: center;
        padding: 0.38rem 0.82rem;
        margin: 0 auto 1rem auto;
        border-radius: 999px;
        background: rgba(24, 6, 12, 0.62);
        border: 1px solid rgba(200, 50, 80, 0.32);
        color: #fce8ee;
        font-size: 0.68rem;
        font-weight: 800;
        letter-spacing: 0.14em;
        text-transform: uppercase;
        box-shadow: inset 0 1px 0 rgba(255,255,255,0.07);
    }

    .guide-step-shell strong {
        position: relative;
        z-index: 1;
        display: block;
        color: #f8fafc;
        font-size: 1.35rem;
        line-height: 1.14;
        margin: 0 auto;
        max-width: 12ch;
        text-shadow: 0 8px 24px rgba(0, 0, 0, 0.32);
    }

    .guide-step-body {
        position: relative;
        z-index: 1;
        max-height: 0;
        overflow: hidden;
        opacity: 0;
        transform: translateY(10px);
        transition: max-height 0.28s ease, opacity 0.24s ease, transform 0.24s ease, margin-top 0.24s ease;
        color: #ece8f8;
        font-size: 0.94rem;
        line-height: 1.5;
        margin: 0 auto;
        padding-top: 0;
        max-width: 24ch;
        font-weight: 600;
    }

    .guide-step-shell:hover .guide-step-body {
        max-height: 240px;
        opacity: 1;
        transform: translateY(0);
        margin-top: 1rem;
        padding-top: 0;
        border-top: none;
    }

    .guide-step-shell:hover .guide-step-pill {
        border-color: rgba(255, 160, 170, 0.38);
        box-shadow: 0 12px 26px rgba(160, 20, 50, 0.22);
    }

    .guide-step-shell:hover {
        transform: translateY(-3px);
        box-shadow: inset 0 1px 0 rgba(255,255,255,0.05), 0 26px 48px rgba(10, 0, 4, 0.38), 0 0 0 1px rgba(200, 50, 80, 0.14);
    }

    .guide-step-shell:hover::before {
        opacity: 0.44;
    }

    .guide-step-shell:hover::after {
        opacity: 0.9;
    }

    /* ── Per-card accent colours ── */

    /* Card 1 — Wine/Burgundy (Source) */
    .guide-step-shell.card-wine {
        background:
            radial-gradient(circle at 50% -8%, rgba(255,255,255,0.06), transparent 26%),
            radial-gradient(circle at 50% 110%, rgba(160, 0, 40, 0.09), transparent 28%),
            linear-gradient(180deg, rgba(14, 6, 10, 0.99), rgba(8, 3, 6, 1));
        border-color: rgba(180, 40, 70, 0.28);
    }
    .guide-step-shell.card-wine::before {
        background: radial-gradient(circle, rgba(200, 30, 60, 0.55) 0.7px, transparent 0.9px);
        background-size: 5px 5px;
    }
    .guide-step-shell.card-wine::after {
        background:
            radial-gradient(circle at 50% 8%, rgba(255,255,255,0.12), transparent 18%),
            radial-gradient(circle at 15% 12%, rgba(180, 30, 60, 0.18), transparent 16%),
            radial-gradient(circle at 85% 18%, rgba(180, 30, 60, 0.18), transparent 18%),
            linear-gradient(180deg, rgba(255,255,255,0.02), transparent 30%, rgba(160, 10, 30, 0.05) 100%);
    }
    .guide-step-shell.card-wine .guide-step-pill {
        border-color: rgba(200, 50, 80, 0.32);
        color: #fce8ee;
    }
    .guide-step-shell.card-wine:hover {
        box-shadow: inset 0 1px 0 rgba(255,255,255,0.05), 0 26px 48px rgba(10, 0, 4, 0.38), 0 0 0 1px rgba(200, 50, 80, 0.14);
    }

    /* Card 2 — Teal/Ocean (Indexing) */
    .guide-step-shell.card-teal {
        background:
            radial-gradient(circle at 50% -8%, rgba(255,255,255,0.06), transparent 26%),
            radial-gradient(circle at 50% 110%, rgba(0, 120, 130, 0.1), transparent 28%),
            linear-gradient(180deg, rgba(4, 14, 16, 0.99), rgba(2, 9, 10, 1));
        border-color: rgba(30, 160, 170, 0.28);
    }
    .guide-step-shell.card-teal::before {
        background: radial-gradient(circle, rgba(20, 200, 200, 0.5) 0.7px, transparent 0.9px);
        background-size: 5px 5px;
    }
    .guide-step-shell.card-teal::after {
        background:
            radial-gradient(circle at 50% 8%, rgba(255,255,255,0.12), transparent 18%),
            radial-gradient(circle at 15% 12%, rgba(20, 180, 180, 0.16), transparent 16%),
            radial-gradient(circle at 85% 18%, rgba(20, 180, 180, 0.16), transparent 18%),
            linear-gradient(180deg, rgba(255,255,255,0.02), transparent 30%, rgba(0, 160, 160, 0.04) 100%);
    }
    .guide-step-shell.card-teal .guide-step-pill {
        border-color: rgba(20, 180, 180, 0.32);
        color: #e0fdfd;
    }
    .guide-step-shell.card-teal:hover {
        box-shadow: inset 0 1px 0 rgba(255,255,255,0.05), 0 26px 48px rgba(0, 6, 8, 0.38), 0 0 0 1px rgba(20, 180, 180, 0.14);
    }

    /* Card 3 — Amber/Gold (Insights) */
    .guide-step-shell.card-amber {
        background:
            radial-gradient(circle at 50% -8%, rgba(255,255,255,0.06), transparent 26%),
            radial-gradient(circle at 50% 110%, rgba(160, 100, 0, 0.09), transparent 28%),
            linear-gradient(180deg, rgba(14, 11, 2, 0.99), rgba(9, 7, 1, 1));
        border-color: rgba(200, 140, 20, 0.28);
    }
    .guide-step-shell.card-amber::before {
        background: radial-gradient(circle, rgba(230, 170, 20, 0.5) 0.7px, transparent 0.9px);
        background-size: 5px 5px;
    }
    .guide-step-shell.card-amber::after {
        background:
            radial-gradient(circle at 50% 8%, rgba(255,255,255,0.12), transparent 18%),
            radial-gradient(circle at 15% 12%, rgba(220, 160, 20, 0.16), transparent 16%),
            radial-gradient(circle at 85% 18%, rgba(220, 160, 20, 0.16), transparent 18%),
            linear-gradient(180deg, rgba(255,255,255,0.02), transparent 30%, rgba(180, 130, 0, 0.04) 100%);
    }
    .guide-step-shell.card-amber .guide-step-pill {
        border-color: rgba(220, 160, 20, 0.32);
        color: #fef9e0;
    }
    .guide-step-shell.card-amber:hover {
        box-shadow: inset 0 1px 0 rgba(255,255,255,0.05), 0 26px 48px rgba(8, 6, 0, 0.38), 0 0 0 1px rgba(220, 160, 20, 0.14);
    }

    @media (max-width: 900px) {
        .guide-step-shell strong {
            max-width: none;
            font-size: 1.18rem;
        }

        .guide-step-body {
            max-width: none;
            font-size: 0.88rem;
        }
    }

    .guide-capability {
        padding: 0.9rem;
        border-radius: 14px;
        background: rgba(255,255,255,0.04);
        border: 1px solid rgba(255,255,255,0.08);
        border-left: 3px solid #06b6d4;
    }

    .guide-capability .mode-title {
        display: block;
        margin-bottom: 4px;
        color: #f8fafc !important;
    }

    .guide-capability .mode-desc {
        color: #94a3b8;
        line-height: 1.6;
    }

    .guide-capability.cap-pdf {
        background: linear-gradient(145deg, rgba(65, 40, 142, 0.22), rgba(20, 24, 46, 0.7));
        border-left-color: #22d3ee;
        border-color: rgba(34, 211, 238, 0.26);
    }

    .guide-capability.cap-pdf .mode-title {
        color: #bae6fd !important;
    }

    .guide-capability.cap-word {
        background: linear-gradient(145deg, rgba(26, 92, 170, 0.24), rgba(18, 28, 52, 0.72));
        border-left-color: #60a5fa;
        border-color: rgba(96, 165, 250, 0.26);
    }

    .guide-capability.cap-word .mode-title {
        color: #bfdbfe !important;
    }

    .guide-capability.cap-text {
        background: linear-gradient(145deg, rgba(80, 46, 170, 0.22), rgba(22, 25, 48, 0.7));
        border-left-color: #a78bfa;
        border-color: rgba(167, 139, 250, 0.26);
    }

    .guide-capability.cap-text .mode-title {
        color: #ddd6fe !important;
    }

    .guide-capability.cap-json {
        background: linear-gradient(145deg, rgba(6, 95, 70, 0.22), rgba(15, 32, 40, 0.72));
        border-left-color: #2dd4bf;
        border-color: rgba(45, 212, 191, 0.26);
    }

    .guide-capability.cap-json .mode-title {
        color: #99f6e4 !important;
    }

    .guide-capability.cap-image {
        background: linear-gradient(145deg, rgba(109, 40, 217, 0.2), rgba(31, 24, 46, 0.72));
        border-left-color: #c084fc;
        border-color: rgba(192, 132, 252, 0.26);
    }

    .guide-capability.cap-image .mode-title {
        color: #e9d5ff !important;
    }

    .guide-capability.cap-csv .mode-title {
        color: #86efac !important;
    }

    .guide-capability.cap-url .mode-title {
        color: #93c5fd !important;
    }

    .guide-panel {
        padding: 0.95rem;
        border-radius: 14px;
        background: rgba(255,255,255,0.04);
        border: 1px solid rgba(255,255,255,0.08);
    }

    .guide-chip-row {
        display: flex;
        flex-wrap: wrap;
        gap: 0.45rem;
        margin-top: 0.8rem;
    }

    .guide-chip {
        display: inline-flex;
        align-items: center;
        padding: 0.32rem 0.6rem;
        border-radius: 999px;
        background: rgba(255,255,255,0.06);
        border: 1px solid rgba(255,255,255,0.08);
        color: #cbd5e1;
        font-size: 0.74rem;
        font-weight: 700;
    }

    .guide-cta {
        padding: 1rem;
        border-radius: 14px;
        background: linear-gradient(135deg, rgba(16,185,129,0.12), rgba(6,182,212,0.08));
        border: 1px solid rgba(52,211,153,0.28);
        margin-bottom: 1.5rem;
    }

    [data-testid="stSidebar"] [data-testid="stExpander"] {
        transition: transform 0.22s ease, box-shadow 0.22s ease, border-color 0.22s ease;
    }

    [data-testid="stSidebar"] [data-testid="stExpander"]:hover {
        transform: translateY(-1px);
        box-shadow: 0 16px 32px rgba(2, 6, 23, 0.32) !important;
        border-color: rgba(125, 211, 252, 0.24) !important;
    }

    [data-testid="stSidebar"] .stSlider,
    [data-testid="stSidebar"] .stSelectbox,
    [data-testid="stSidebar"] .stMultiSelect,
    [data-testid="stSidebar"] .stTextInput,
    [data-testid="stSidebar"] .stButton,
    [data-testid="stSidebar"] .stDownloadButton {
        padding: 0.35rem 0.25rem;
        border-radius: 14px;
        transition: background 0.22s ease, box-shadow 0.22s ease;
    }

    [data-testid="stSidebar"] .stSlider:hover,
    [data-testid="stSidebar"] .stSelectbox:hover,
    [data-testid="stSidebar"] .stMultiSelect:hover,
    [data-testid="stSidebar"] .stTextInput:hover,
    [data-testid="stSidebar"] .stButton:hover,
    [data-testid="stSidebar"] .stDownloadButton:hover {
        background: rgba(255, 255, 255, 0.035);
        box-shadow: inset 0 0 0 1px rgba(255, 255, 255, 0.04);
    }

    [data-testid="stSidebar"] .stMultiSelect [data-baseweb="tag"] {
        background: linear-gradient(135deg, rgba(14, 116, 144, 0.9), rgba(30, 64, 175, 0.9)) !important;
        color: #e0f2fe !important;
        border: 1px solid rgba(125, 211, 252, 0.45) !important;
    }

    [data-testid="stSidebar"] .stMultiSelect [data-baseweb="tag"] svg {
        color: #dbeafe !important;
    }

    html[data-theme="light"] [data-testid="stSidebar"] .stMultiSelect [data-baseweb="tag"] {
        background: linear-gradient(135deg, rgba(59, 130, 246, 0.9), rgba(14, 165, 233, 0.9)) !important;
        color: #f8fafc !important;
        border: 1px solid rgba(59, 130, 246, 0.45) !important;
    }

    .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
        max-width: 1360px;
    }

    .header-container {
        position: relative;
        overflow: hidden;
        isolation: isolate;
        padding: 2.25rem;
        margin-bottom: 1.75rem;
        border-radius: var(--radius-xl);
        border: 1px solid rgba(255, 255, 255, 0.1);
        background:
            radial-gradient(circle at 14% 18%, rgba(233, 213, 255, 0.12), transparent 20%),
            radial-gradient(circle at 86% 14%, rgba(139, 92, 246, 0.24), transparent 24%),
            linear-gradient(135deg, rgba(7, 10, 31, 0.98) 0%, rgba(8, 13, 39, 0.99) 42%, rgba(3, 5, 18, 0.99) 100%);
        box-shadow: var(--shadow-xl);
    }

    .header-container::before {
        content: "";
        position: absolute;
        top: -70px;
        right: -70px;
        width: 450px;
        height: 450px;
        background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 620 620'%3E%3Cdefs%3E%3ClinearGradient id='cy' x1='0' y1='0' x2='1' y2='1'%3E%3Cstop offset='0%25' stop-color='%238be9ff'/%3E%3Cstop offset='100%25' stop-color='%2360a5fa'/%3E%3C/linearGradient%3E%3C/defs%3E%3Cg fill='none' stroke-linecap='round' stroke-linejoin='round'%3E%3Cg stroke='url(%23cy)' stroke-width='4.1' opacity='0.78'%3E%3Cpath d='M422 86c73 24 127 88 138 167 14 102-44 201-138 243'/%3E%3Cpath d='M396 104c58 22 99 73 108 135 12 81-29 161-102 198'/%3E%3Cpath d='M369 123c44 19 74 59 82 108 10 64-20 127-74 160'/%3E%3C/g%3E%3Cg stroke='%237dd3fc' stroke-width='4.4' opacity='0.92'%3E%3Cpath d='M360 134c-49 16-82 48-92 92-8 33 0 67 23 93-16 18-22 39-18 63 5 26 23 47 51 60-23 19-37 44-38 73-1 34 14 65 43 87'/%3E%3Cpath d='M302 205c13 2 24 10 29 21'/%3E%3Cpath d='M331 228c26 3 46 14 61 33'/%3E%3Cpath d='M316 271c-11 16-15 34-12 53 29 7 52 20 69 40'/%3E%3Cpath d='M311 338c13-2 25 0 36 7'/%3E%3Cpath d='M308 365c13 10 27 13 42 10'/%3E%3Cpath d='M309 404c20 2 38 10 54 23'/%3E%3C/g%3E%3Cg stroke='%2367e8f9' stroke-width='2.5' opacity='0.6'%3E%3Cpath d='M156 180l122 22 99-30 155 38'/%3E%3Cpath d='M147 226l128 15 108-20 157 43'/%3E%3Cpath d='M144 274l131 12 112-10 156 49'/%3E%3Cpath d='M148 322l129 9 113 2 149 54'/%3E%3Cpath d='M161 370l118 12 109 12 136 57'/%3E%3Cpath d='M178 414l102 17 100 22 122 60'/%3E%3C/g%3E%3Cg stroke='%2393c5fd' stroke-width='1.9' opacity='0.58'%3E%3Cpath d='M187 195l57 33 47-31 45 36 52-25'/%3E%3Cpath d='M178 244l64 24 43-24 52 31 48-19'/%3E%3Cpath d='M175 289l68 18 44-16 54 26 45-14'/%3E%3Cpath d='M180 333l66 14 47-10 53 23 40-12'/%3E%3Cpath d='M190 376l61 13 49-5 51 20 34-8'/%3E%3Cpath d='M204 418l54 14 50 2 47 17 28-3'/%3E%3C/g%3E%3Cg stroke='%2367e8f9' stroke-width='2.9' opacity='0.72'%3E%3Cpath d='M118 148h22m-11-11v22'/%3E%3Cpath d='M134 188h22m-11-11v22'/%3E%3Cpath d='M148 228h22m-11-11v22'/%3E%3Cpath d='M159 268h22m-11-11v22'/%3E%3Cpath d='M170 308h22m-11-11v22'/%3E%3Cpath d='M181 348h22m-11-11v22'/%3E%3Cpath d='M191 388h22m-11-11v22'/%3E%3C/g%3E%3Cg fill='%238be9ff' stroke='none' opacity='0.88'%3E%3Ccircle cx='360' cy='134' r='4.8'/%3E%3Ccircle cx='302' cy='205' r='4.1'/%3E%3Ccircle cx='331' cy='228' r='4.1'/%3E%3Ccircle cx='316' cy='271' r='4.1'/%3E%3Ccircle cx='311' cy='338' r='4.1'/%3E%3Ccircle cx='308' cy='365' r='4.1'/%3E%3Ccircle cx='309' cy='404' r='4.1'/%3E%3Ccircle cx='336' cy='497' r='4.9'/%3E%3Ccircle cx='389' cy='173' r='3.5'/%3E%3Ccircle cx='385' cy='221' r='3.5'/%3E%3Ccircle cx='387' cy='278' r='3.5'/%3E%3Ccircle cx='386' cy='338' r='3.5'/%3E%3Ccircle cx='381' cy='395' r='3.5'/%3E%3C/g%3E%3Cg stroke='%2367e8f9' stroke-width='3.2' opacity='0.42'%3E%3Cpath d='M425 86l58-54M467 104l68-39M504 130l74-22'/%3E%3C/g%3E%3C/g%3E%3C/svg%3E");
        background-repeat: no-repeat;
        background-size: contain;
        opacity: 0.38;
        pointer-events: none;
        z-index: 0;
        filter: drop-shadow(0 0 22px rgba(103, 232, 249, 0.28));
        transform: rotate(-3deg);
        animation: heroAiSigPulse 7s ease-in-out infinite;
    }

    @keyframes heroAiSigPulse {
        0%,
        100% {
            opacity: 0.32;
            transform: rotate(-3deg) translateY(0px) scale(1);
        }
        50% {
            opacity: 0.45;
            transform: rotate(-1deg) translateY(-8px) scale(1.05);
        }
    }

    .header-container::after {
        content: "";
        position: absolute;
        inset: auto -8% -22% auto;
        width: 240px;
        height: 240px;
        border-radius: 999px;
        background: radial-gradient(circle, rgba(208, 179, 255, 0.22), transparent 68%);
        pointer-events: none;
        z-index: 0;
    }

    .header-container > * {
        position: relative;
        z-index: 1;
    }

    .header-container h1 {
        margin: 0;
        font-family: 'Sora', 'Outfit', sans-serif;
        font-size: clamp(2.6rem, 5vw, 4.4rem);
        line-height: 0.98;
        font-weight: 700;
        letter-spacing: -0.04em;
        color: #ffffff;
    }

    /* hide Streamlit's injected anchor icon on our custom hero heading */
    .header-container h1 a,
    .header-container h1 a:hover {
        display: none !important;
    }

    .hero-accent {
        display: inline-block;
        background: linear-gradient(135deg, #f5f3ff 0%, #d8b4fe 38%, #c4b5fd 74%, #93c5fd 100%);
        -webkit-background-clip: text;
        background-clip: text;
        color: transparent;
        -webkit-text-fill-color: transparent;
    }

    .header-container p {
        max-width: 860px;
        margin: 1rem 0 1.4rem 0;
        font-size: 1.05rem;
        line-height: 1.8;
        color: var(--text-1);
    }

    .section-heading {
        display: flex;
        align-items: center;
        gap: 0.9rem;
        margin: 0 0 1.1rem 0;
    }

    .section-heading i {
        width: 2.6rem;
        height: 2.6rem;
        display: inline-flex;
        align-items: center;
        justify-content: center;
        border-radius: 14px;
        background: linear-gradient(135deg, rgba(124, 58, 237, 0.22), rgba(59, 130, 246, 0.16));
        border: 1px solid rgba(196, 181, 253, 0.16);
        color: #efe9ff;
        box-shadow: inset 0 1px 0 rgba(255,255,255,0.06);
        font-size: 1.05rem;
    }

    .section-heading h3 {
        margin: 0 !important;
        font-size: clamp(1.6rem, 2vw, 2.15rem);
        line-height: 1.05;
    }

    .guide-section-heading {
        display: flex;
        align-items: center;
        gap: 0.55rem;
        margin: 0.2rem 0 0.7rem 0;
    }

    .guide-section-heading i {
        width: 1.55rem;
        height: 1.55rem;
        display: inline-flex;
        align-items: center;
        justify-content: center;
        border-radius: 10px;
        border: 1px solid rgba(196, 181, 253, 0.18);
        background: linear-gradient(135deg, rgba(124, 58, 237, 0.2), rgba(6, 182, 212, 0.16));
        color: #efe9ff;
        font-size: 0.78rem;
        flex-shrink: 0;
    }

    .guide-section-heading h4 {
        margin: 0 !important;
        color: #f8fafc;
        font-size: 1.08rem;
        font-weight: 700;
        letter-spacing: 0.01em;
    }

    .hero-stats {
        display: grid;
        grid-template-columns: repeat(3, minmax(0, 1fr));
        gap: 1rem;
        margin-top: 1.25rem;
    }

    .hero-stat {
        padding: 1rem 1.1rem;
        border-radius: 18px;
        background: linear-gradient(145deg, rgba(16, 19, 46, 0.82), rgba(10, 14, 32, 0.74));
        border: 1px solid rgba(196, 181, 253, 0.16);
        backdrop-filter: blur(12px);
        position: relative;
        overflow: hidden;
        isolation: isolate;
        animation: heroCardFloat 5.2s ease-in-out infinite;
        transition: transform 260ms ease, box-shadow 260ms ease, border-color 260ms ease;
        will-change: transform;
    }

    .hero-stat:hover {
        transform: translateY(-6px) scale(1.01);
        box-shadow: 0 18px 36px rgba(2, 6, 23, 0.34), inset 0 1px 0 rgba(255, 255, 255, 0.06);
    }

    .hero-stats .hero-stat:nth-child(1) {
        background: linear-gradient(145deg, rgba(56, 30, 114, 0.34), rgba(12, 16, 39, 0.82));
        border-color: rgba(196, 181, 253, 0.34);
        box-shadow: inset 0 1px 0 rgba(196, 181, 253, 0.14);
        animation-delay: 0s;
    }

    .hero-stats .hero-stat:nth-child(2) {
        background: linear-gradient(145deg, rgba(8, 65, 88, 0.34), rgba(10, 17, 35, 0.82));
        border-color: rgba(103, 232, 249, 0.33);
        box-shadow: inset 0 1px 0 rgba(103, 232, 249, 0.12);
        animation-delay: 0.9s;
    }

    .hero-stats .hero-stat:nth-child(3) {
        background: linear-gradient(145deg, rgba(95, 30, 70, 0.34), rgba(14, 15, 34, 0.82));
        border-color: rgba(249, 168, 212, 0.32);
        box-shadow: inset 0 1px 0 rgba(249, 168, 212, 0.12);
        animation-delay: 1.8s;
    }

    @keyframes heroCardFloat {
        0%,
        100% {
            transform: translate3d(0px, 0px, 0px);
        }
        50% {
            transform: translate3d(4px, -11px, 0px);
        }
    }

    @media (prefers-reduced-motion: reduce) {
        .hero-stat {
            animation: none !important;
            transition: none !important;
        }
        .hero-stat:hover {
            transform: none !important;
        }
    }

    .hero-stat::before {
        content: "";
        position: absolute;
        inset: 0;
        background-repeat: no-repeat;
        background-position: right -22px bottom -24px;
        background-size: 66% auto;
        opacity: 0.22;
        pointer-events: none;
        z-index: 0;
    }

    .hero-stat > * {
        position: relative;
        z-index: 1;
    }

    .hero-stats .hero-stat:nth-child(1)::before {
        background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 320 220'%3E%3Cg fill='none' stroke='%23c4b5fd' stroke-width='3'%3E%3Crect x='134' y='28' width='128' height='164' rx='14'/%3E%3Cpath d='M214 28v46h48'/%3E%3Crect x='82' y='48' width='120' height='152' rx='14' opacity='0.62'/%3E%3Crect x='42' y='74' width='110' height='130' rx='12' opacity='0.45'/%3E%3C/g%3E%3C/svg%3E");
    }

    .hero-stats .hero-stat:nth-child(2)::before {
        background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 320 220'%3E%3Cg fill='none' stroke='%2367e8f9' stroke-width='3'%3E%3Cpath d='M38 42h184a16 16 0 0 1 16 16v64a16 16 0 0 1-16 16H118l-44 34v-34H38a16 16 0 0 1-16-16V58a16 16 0 0 1 16-16z'/%3E%3Cpath d='M126 104h106a14 14 0 0 1 14 14v44a14 14 0 0 1-14 14h-62l-28 24v-24h-16a14 14 0 0 1-14-14v-44a14 14 0 0 1 14-14z' opacity='0.55'/%3E%3C/g%3E%3C/svg%3E");
    }

    .hero-stats .hero-stat:nth-child(3)::before {
        background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 320 220'%3E%3Cg fill='none' stroke='%23f9a8d4' stroke-width='3'%3E%3Cpath d='M22 174h276'/%3E%3Cpath d='M44 156l54-42 46 22 58-64 58 24' stroke-linecap='round' stroke-linejoin='round'/%3E%3Ccircle cx='44' cy='156' r='6' fill='%23f9a8d4'/%3E%3Ccircle cx='98' cy='114' r='6' fill='%23f9a8d4'/%3E%3Ccircle cx='144' cy='136' r='6' fill='%23f9a8d4'/%3E%3Ccircle cx='202' cy='72' r='6' fill='%23f9a8d4'/%3E%3Ccircle cx='260' cy='96' r='6' fill='%23f9a8d4'/%3E%3C/g%3E%3C/svg%3E");
    }

    .header-container .hero-stat strong {
        display: block;
        color: #dbeafe !important;
        font-size: 1rem;
        font-weight: 700;
        margin-bottom: 0.2rem;
    }

    .header-container .hero-stats .hero-stat:nth-child(1) strong {
        color: #c4b5fd !important;
    }

    .header-container .hero-stats .hero-stat:nth-child(2) strong {
        color: #67e8f9 !important;
    }

    .header-container .hero-stats .hero-stat:nth-child(3) strong {
        color: #f9a8d4 !important;
    }

    .hero-stat span {
        color: var(--text-2);
        font-size: 0.9rem;
    }

    .chat-container,
    .metric-card,
    [data-testid="stExpander"],
    [data-testid="stFileUploader"] section,
    div[data-baseweb="select"] > div,
    .stTextInput > div > div > input,
    .stTextArea textarea {
        background: var(--panel) !important;
        border: 1px solid var(--border) !important;
        box-shadow: var(--shadow-md) !important;
        color: var(--text-0) !important;
        backdrop-filter: blur(14px);
    }

    [data-testid="stFileUploader"] section,
    [data-testid="stFileUploaderDropzone"] {
        background-color: var(--panel) !important;
        color: var(--text-1) !important;
    }

    [data-testid="stFileUploader"] section {
        overflow: hidden !important;
    }

    html[data-theme="light"] [data-testid="stFileUploader"] section,
    html[data-theme="light"] [data-testid="stFileUploaderDropzone"] {
        background-color: rgba(255, 255, 255, 0.86) !important;
        color: #334155 !important;
    }

    [data-testid="stFileUploaderDropzone"] [data-testid="stMarkdownContainer"],
    [data-testid="stFileUploaderDropzone"] [data-testid="stMarkdownContainer"] > div,
    [data-testid="stFileUploaderDropzone"] [data-testid="stMarkdownContainer"] > div > div {
        background: transparent !important;
        box-shadow: none !important;
        border: none !important;
    }

    .chat-container {
        border-radius: var(--radius-lg);
        padding: 1.35rem;
        margin-bottom: 1.25rem;
        min-height: 120px;
        max-height: 540px;
        overflow-y: auto;
    }

    .chat-container::-webkit-scrollbar {
        width: 10px;
    }

    .chat-container::-webkit-scrollbar-thumb {
        background: linear-gradient(180deg, rgba(79, 70, 229, 0.9), rgba(6, 182, 212, 0.9));
        border-radius: 999px;
    }

    .message-user,
    .message-assistant {
        position: relative;
        border-radius: 18px;
        padding: 1rem 1.1rem;
        margin: 0.9rem 0;
        max-width: 78%;
        word-wrap: break-word;
        line-height: 1.75;
        border: 1px solid rgba(255, 255, 255, 0.08);
    }

    .message-user {
        background: linear-gradient(135deg, rgba(79, 70, 229, 0.92) 0%, rgba(29, 78, 216, 0.88) 45%, rgba(6, 182, 212, 0.84) 100%);
        color: #f8fafc;
        margin-left: auto;
        box-shadow: 0 18px 36px rgba(37, 99, 235, 0.24);
    }

    .message-assistant {
        background: linear-gradient(135deg, rgba(15, 23, 42, 0.96) 0%, rgba(17, 24, 39, 0.98) 100%);
        color: var(--text-1);
        border-left: 4px solid var(--brand-3);
        box-shadow: 0 16px 34px rgba(15, 23, 42, 0.35);
    }

    .metric-card {
        border-radius: 18px;
        padding: 1.15rem;
        text-align: center;
        color: var(--text-1);
        border-top: 1px solid rgba(6, 182, 212, 0.45);
    }

    .metric-card strong {
        color: var(--text-0);
    }

    /* ── animated rotating glow border ── */
    @property --btn-angle {
        syntax: '<angle>';
        initial-value: 0deg;
        inherits: false;
    }

    @keyframes btnBorderSpin {
        to { --btn-angle: 360deg; }
    }

    .stButton > button,
    .stDownloadButton > button {
        position: relative;
        min-height: 3.1rem;
        border-radius: 14px;
        background: rgba(4, 6, 20, 0.88);
        backdrop-filter: blur(12px);
        color: #f0f4ff;
        font-weight: 600;
        letter-spacing: 0.02em;
        border: none;
        /* spinning conic-gradient border via outline-like box-shadow + pseudo layer */
        box-shadow:
            0 0 0 1.5px rgba(139, 92, 246, 0.18),
            inset 0 1px 0 rgba(255, 255, 255, 0.06);
        transition: transform 0.2s ease, background 0.2s ease;
        overflow: hidden;
        z-index: 0;
    }

    /* spinning glow layer via ::before */
    .stButton > button::before,
    .stDownloadButton > button::before {
        content: "";
        position: absolute;
        inset: -2px;
        border-radius: inherit;
        background: conic-gradient(
            from var(--btn-angle),
            transparent 0deg,
            rgba(139, 92, 246, 0.0) 40deg,
            rgba(167, 139, 250, 0.9) 80deg,
            rgba(99, 179, 237, 0.7) 110deg,
            rgba(139, 92, 246, 0.0) 150deg,
            transparent 360deg
        );
        animation: btnBorderSpin 2.4s linear infinite;
        z-index: -1;
        border-radius: 15px;
    }

    /* inner masking layer to keep the button background solid */
    .stButton > button::after,
    .stDownloadButton > button::after {
        content: "";
        position: absolute;
        inset: 1.5px;
        border-radius: 12px;
        background: rgba(4, 6, 20, 0.92);
        z-index: -1;
    }

    .stButton > button:hover,
    .stDownloadButton > button:hover {
        background: rgba(9, 12, 30, 0.92);
        transform: translateY(-1px);
        box-shadow: 0 0 0 1.5px rgba(139, 92, 246, 0.12), 0 8px 24px rgba(139,92,246,0.18), inset 0 1px 0 rgba(255,255,255,0.09);
    }

    .stButton > button:active,
    .stDownloadButton > button:active {
        background: rgba(7, 10, 24, 0.92);
        transform: translateY(0);
    }

    .stButton > button:disabled,
    .stDownloadButton > button:disabled {
        background: rgba(15, 15, 30, 0.5);
        color: rgba(248, 250, 252, 0.38);
        box-shadow: 0 0 0 1.5px rgba(100, 100, 140, 0.18);
        border: none;
    }

    .stButton > button:disabled::before,
    .stDownloadButton > button:disabled::before {
        display: none;
    }

    /* file uploader browse button — same spinning glow */
    [data-testid="stFileUploader"] button,
    [data-testid="stFileUploaderDropzone"] button {
        position: relative !important;
        border-radius: 12px !important;
        background: rgba(4, 6, 20, 0.0) !important;
        color: #f0f4ff !important;
        font-weight: 600 !important;
        border: none !important;
        box-shadow: 0 0 0 1.5px rgba(139, 92, 246, 0.18) !important;
        overflow: hidden !important;
        z-index: 0 !important;
        transition: transform 0.2s ease !important;
    }

    [data-testid="stFileUploader"] button::before,
    [data-testid="stFileUploaderDropzone"] button::before {
        content: "" !important;
        position: absolute !important;
        inset: -2px !important;
        border-radius: 13px !important;
        background: conic-gradient(
            from var(--btn-angle),
            transparent 0deg,
            rgba(139, 92, 246, 0.0) 40deg,
            rgba(167, 139, 250, 0.9) 80deg,
            rgba(99, 179, 237, 0.7) 110deg,
            rgba(139, 92, 246, 0.0) 150deg,
            transparent 360deg
        ) !important;
        animation: btnBorderSpin 2.4s linear infinite !important;
        z-index: -1 !important;
    }

    [data-testid="stFileUploader"] button::after,
    [data-testid="stFileUploaderDropzone"] button::after {
        content: "" !important;
        position: absolute !important;
        inset: 1.5px !important;
        border-radius: 10px !important;
        background: rgba(4, 6, 20, 0.92) !important;
        z-index: -1 !important;
    }

    [data-testid="stFileUploader"] button:hover,
    [data-testid="stFileUploaderDropzone"] button:hover {
        transform: translateY(-1px) !important;
        box-shadow: 0 8px 22px rgba(139, 92, 246, 0.18) !important;
    }

    .floating-backdrop {
        position: fixed;
        inset: 0;
        z-index: 9998;
        background: rgba(2, 6, 23, 0.45);
        backdrop-filter: blur(2px);
        pointer-events: none;
    }

    .floating-modal {
        position: fixed;
        top: 7.2rem;
        left: 50%;
        transform: translateX(-50%);
        width: min(760px, 86vw);
        z-index: 9999;
        padding: 1.05rem 1.2rem;
        border-radius: 18px;
        border: 1px solid rgba(125, 211, 252, 0.32);
        background: linear-gradient(140deg, rgba(15, 23, 42, 0.94) 0%, rgba(17, 24, 39, 0.96) 55%, rgba(8, 145, 178, 0.24) 100%);
        box-shadow: 0 30px 60px rgba(2, 6, 23, 0.55);
        backdrop-filter: blur(14px);
        pointer-events: auto;
        user-select: none;
    }

    .floating-modal-model {
        min-height: 330px;
        padding-bottom: 6rem;
    }

    .floating-window-bar {
        display: flex;
        align-items: center;
        justify-content: space-between;
        gap: 0.55rem;
        margin: -0.2rem -0.2rem 0.55rem -0.2rem;
        padding: 0.35rem 0.45rem;
        border-radius: 11px;
        border: 1px solid rgba(255,255,255,0.08);
        background: rgba(255,255,255,0.04);
        cursor: grab;
    }

    .floating-window-title {
        color: #f8fafc;
        font-size: 0.9rem;
        font-weight: 800;
        letter-spacing: 0.01em;
    }

    .floating-window-actions {
        display: flex;
        align-items: center;
        gap: 0.35rem;
    }

    .floating-window-btn {
        width: 1.5rem;
        height: 1.5rem;
        border-radius: 999px;
        border: 1px solid rgba(186, 230, 253, 0.28);
        background: rgba(255,255,255,0.08);
        color: #e2e8f0;
        font-size: 0.86rem;
        line-height: 1;
        display: inline-flex;
        align-items: center;
        justify-content: center;
        cursor: pointer;
    }

    .floating-window-btn:hover {
        background: rgba(186, 230, 253, 0.2);
        color: #ffffff;
    }

    .floating-modal h4 {
        margin: 0 0 0.35rem 0;
        color: #f8fafc;
    }

    .floating-modal p,
    .floating-modal li {
        color: #dbe4f3;
        font-size: 0.92rem;
        line-height: 1.7;
    }

    .hidden-close-action,
    .hidden-action-bank {
        position: fixed;
        left: -9999px;
        top: -9999px;
        width: 1px;
        height: 1px;
        overflow: hidden;
        opacity: 0;
        pointer-events: none;
    }

    .floating-quick-actions {
        display: grid;
        grid-template-columns: repeat(3, minmax(0, 1fr));
        gap: 0.75rem;
        margin-top: 1.2rem;
    }

    .floating-quick-actions.two-up {
        grid-template-columns: 1fr;
    }

    .floating-quick-btn {
        min-height: 3rem;
        border-radius: 14px;
        border: 1px solid rgba(167, 139, 250, 0.26);
        background: rgba(6, 10, 28, 0.72);
        color: #f0f4ff;
        font-weight: 700;
        letter-spacing: 0.01em;
        box-shadow: 0 0 0 1px rgba(139, 92, 246, 0.12), inset 0 1px 0 rgba(255, 255, 255, 0.05);
        cursor: pointer;
        transition: transform 0.18s ease, box-shadow 0.18s ease, border-color 0.18s ease;
    }

    .floating-quick-btn:hover {
        transform: translateY(-1px);
        border-color: rgba(196, 181, 253, 0.42);
        box-shadow: 0 12px 24px rgba(76, 29, 149, 0.22), 0 0 0 1px rgba(139, 92, 246, 0.16);
    }

    @media (max-width: 900px) {
        .floating-quick-actions {
            grid-template-columns: 1fr;
        }
    }

    .stTextInput > div > div > input,
    .stTextArea textarea {
        border-radius: 16px !important;
        padding: 0.95rem 1rem !important;
        color: var(--text-0) !important;
    }

    .stTextInput input:disabled,
    .stTextArea textarea:disabled {
        color: var(--text-2) !important;
        -webkit-text-fill-color: var(--text-2) !important;
        opacity: 1 !important;
        cursor: not-allowed !important;
    }

    .stTextInput > div > div > input::placeholder,
    .stTextArea textarea::placeholder {
        color: var(--text-2) !important;
    }

    .stTextInput > div > div > input:focus,
    .stTextArea textarea:focus {
        border-color: rgba(139, 92, 246, 0.6) !important;
        box-shadow: 0 0 0 0.2rem rgba(139, 92, 246, 0.18), var(--shadow-md) !important;
        outline: none !important;
    }

    /* Override BaseWeb input/textarea wrapper focus border (prevents red outline) */
    [data-baseweb="input"]:focus-within,
    [data-baseweb="textarea"]:focus-within,
    [data-baseweb="base-input"]:focus-within,
    div[data-baseweb="input"] > div,
    div[data-baseweb="textarea"] > div {
        border-color: rgba(139, 92, 246, 0.6) !important;
        box-shadow: 0 0 0 0.2rem rgba(139, 92, 246, 0.18) !important;
        outline: none !important;
    }

    /* Catch any remaining red/error border from Streamlit theme */
    .stTextInput > div,
    .stTextArea > div {
        border-color: rgba(139, 92, 246, 0.6) !important;
    }
    .stTextInput > div > div,
    .stTextArea > div > div {
        border-color: rgba(139, 92, 246, 0.6) !important;
        box-shadow: none !important;
    }

    .stTextInput > div > div:has(input:disabled),
    .stTextArea > div > div:has(textarea:disabled) {
        background: rgba(71, 85, 105, 0.24) !important;
        border-color: rgba(148, 163, 184, 0.38) !important;
        box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.04) !important;
    }

    html[data-theme="light"] .stTextInput > div > div:has(input:disabled),
    html[data-theme="light"] .stTextArea > div > div:has(textarea:disabled) {
        background: rgba(226, 232, 240, 0.9) !important;
        border-color: rgba(148, 163, 184, 0.55) !important;
    }
    .stTextInput > div > div:focus-within,
    .stTextArea > div > div:focus-within {
        border-color: rgba(139, 92, 246, 0.6) !important;
        box-shadow: 0 0 0 0.2rem rgba(139, 92, 246, 0.18) !important;
    }

    .stSelectbox label,
    .stSlider label,
    .stFileUploader label,
    .stTextInput label,
    .stCaption,
    .stMarkdown,
    .stAlert {
        color: var(--text-1);
    }

    .stMarkdown p,
    .stMarkdown li,
    .stMarkdown small {
        color: var(--text-1) !important;
    }

    .stMarkdown h1,
    .stMarkdown h2,
    .stMarkdown h3,
    .stMarkdown h4 {
        background: linear-gradient(135deg, #ffffff 0%, #e9d5ff 30%, #d8b4fe 58%, #93c5fd 100%);
        -webkit-background-clip: text;
        background-clip: text;
        color: transparent !important;
        -webkit-text-fill-color: transparent !important;
    }

    .stMarkdown strong {
        color: var(--text-0) !important;
    }

    .stTabs [data-baseweb="tab-list"] {
        display: grid;
        grid-template-columns: repeat(7, minmax(0, 1fr));
        gap: 0.55rem;
        width: 100%;
        padding: 0.55rem;
        border-radius: 20px;
        background: rgba(255, 255, 255, 0.04);
        border: 1px solid var(--border);
        backdrop-filter: blur(14px);
        box-shadow: var(--shadow-md);
    }

    .stTabs [data-baseweb="tab"] {
        width: 100%;
        min-height: 3rem;
        border-radius: 14px;
        background: transparent;
        color: var(--text-2);
        font-weight: 700;
        border: 1px solid transparent;
        transition: all 0.22s ease;
        justify-content: center;
        padding: 0.7rem 0.55rem;
        white-space: normal !important;
        overflow-wrap: anywhere;
        word-break: break-word;
        text-align: center;
        line-height: 1.15;
        font-size: 0.88rem;
    }

    .stTabs [data-baseweb="tab"]:hover {
        color: var(--text-0);
        background: rgba(255, 255, 255, 0.05);
        border-color: rgba(255, 255, 255, 0.08);
    }

    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, rgba(109, 40, 217, 0.92), rgba(139, 92, 246, 0.82)) !important;
        color: #ffffff !important;
        border-color: rgba(196, 181, 253, 0.18) !important;
        box-shadow: 0 14px 28px rgba(109, 40, 217, 0.28);
    }

    /* ── slider: all track segments ── */
    .stSlider [data-baseweb="slider"] > div > div,
    [data-testid="stSlider"] [data-baseweb="slider"] > div > div {
        background: linear-gradient(90deg, #7c3aed, #a78bfa) !important;
    }

    /* slider unfilled track */
    .stSlider [data-baseweb="slider"] > div:first-child,
    [data-testid="stSlider"] [data-baseweb="slider"] > div:first-child {
        background: rgba(124, 58, 237, 0.18) !important;
    }

    /* slider thumb */
    .stSlider [data-baseweb="slider"] [role="slider"],
    [data-testid="stSlider"] [role="slider"] {
        background: #c4b5fd !important;
        border: 2px solid #7c3aed !important;
        box-shadow: 0 0 0 4px rgba(124, 58, 237, 0.22) !important;
    }

    /* ── toggle / switch — exhaustive overrides ── */
    [data-testid="stToggleSwitch"] input[type="checkbox"] {
        accent-color: #8b5cf6 !important;
    }

    [data-testid="stToggleSwitch"] > label > div,
    [data-baseweb="toggle"],
    [data-baseweb="toggle"] > div,
    [role="switch"],
    [role="switch"] > div {
        background-color: rgba(109, 40, 217, 0.28) !important;
        border-color: rgba(167, 139, 250, 0.28) !important;
    }

    /* track when ON */
    [data-testid="stToggleSwitch"] input:checked ~ div,
    [data-testid="stToggleSwitch"] input[type="checkbox"]:checked + div,
    [data-testid="stToggleSwitch"] label:has(input:checked) > div,
    [role="switch"][aria-checked="true"],
    [role="switch"][aria-checked="true"] > div,
    [data-baseweb="toggle"][aria-checked="true"] > div {
        background: linear-gradient(135deg, #7c3aed, #a78bfa) !important;
        border-color: rgba(196, 181, 253, 0.4) !important;
        box-shadow: 0 0 12px rgba(124, 58, 237, 0.38) !important;
    }

    /* thumb pill */
    [role="switch"] > div > div,
    [role="switch"] span,
    [data-baseweb="toggle"] > div > div {
        background-color: #f3f0ff !important;
        box-shadow: 0 1px 4px rgba(0,0,0,0.22) !important;
    }

    /* ── active tab: kill the red underline ── */
    [data-baseweb="tab-highlight"] {
        background: linear-gradient(90deg, #7c3aed, #a78bfa) !important;
    }

    [data-baseweb="tab-border"] {
        background-color: rgba(124, 58, 237, 0.18) !important;
    }

    /* Suggestion buttons in chat panel */
    [data-testid^="suggest_"] button,
    button[kind="secondary"][data-testid*="suggest_"] {
        background: rgba(124, 58, 237, 0.10) !important;
        border: 1px solid rgba(167, 139, 250, 0.28) !important;
        border-radius: 12px !important;
        color: var(--text-1) !important;
        text-align: left !important;
        font-size: 0.87rem !important;
        padding: 0.55rem 1rem !important;
        transition: background 0.18s, border-color 0.18s;
    }
    [data-testid^="suggest_"] button:hover {
        background: rgba(124, 58, 237, 0.22) !important;
        border-color: rgba(167, 139, 250, 0.55) !important;
    }

    [data-testid="stFileUploader"] section {
        border-radius: 20px !important;
    }

    [data-testid="stFileUploaderDropzone"] {
        background: rgba(255, 255, 255, 0.03) !important;
        border: 1px dashed rgba(167, 139, 250, 0.38) !important;
        border-radius: 18px !important;
    }

    [data-testid="stFileUploaderDropzone"] [data-testid="stMarkdownContainer"] p,
    [data-testid="stFileUploaderDropzone"] [data-testid="stMarkdownContainer"] small,
    [data-testid="stFileUploaderDropzone"] [data-testid="stMarkdownContainer"] span,
    [data-testid="stFileUploader"] small,
    [data-testid="stFileUploader"] span {
        color: var(--text-1) !important;
        opacity: 0.96 !important;
    }

    [data-testid="stFileUploaderDropzone"]:hover {
        border-color: rgba(167, 139, 250, 0.75) !important;
        background: rgba(124, 58, 237, 0.08) !important;
    }

    [data-testid="stExpander"] {
        border-radius: 18px !important;
    }

    [data-testid="stMetric"] {
        background: rgba(255, 255, 255, 0.04);
        border: 1px solid var(--border);
        padding: 0.9rem 1rem;
        border-radius: 16px;
    }

    hr {
        border: 0;
        height: 1px;
        background: linear-gradient(90deg, transparent, rgba(148, 163, 184, 0.35), transparent);
    }

    select {
        cursor: pointer !important;
    }

    @media (max-width: 900px) {
        .stTabs [data-baseweb="tab-list"] {
            grid-template-columns: repeat(3, minmax(0, 1fr));
        }

        .header-container {
            padding: 1.4rem;
        }

        .header-container::before {
            top: -30px;
            right: -30px;
            width: 250px;
            height: 250px;
            opacity: 0.24;
        }

        .hero-stats {
            grid-template-columns: 1fr;
        }

        .message-user,
        .message-assistant {
            max-width: 100%;
        }

        .chat-container {
            height: 460px;
        }
    }

    @media (max-width: 640px) {
        .stTabs [data-baseweb="tab-list"] {
            grid-template-columns: repeat(2, minmax(0, 1fr));
        }

        .stTabs [data-baseweb="tab"] {
            min-height: 3.5rem;
            font-size: 0.82rem;
            line-height: 1.2;
            white-space: normal;
        }

        .header-container::before {
            width: 190px;
            height: 190px;
            top: -24px;
            right: -22px;
            opacity: 0.18;
        }
    }
    </style>
""", unsafe_allow_html=True)

custom_hero_image_uri = _get_custom_hero_image_data_uri()
if custom_hero_image_uri:
    st.markdown(
        f"""
        <style>
        .header-container::before {{
            background-image: url('{custom_hero_image_uri}') !important;
            top: -52px !important;
            right: -58px !important;
            width: 500px !important;
            height: 500px !important;
            opacity: 0.42 !important;
            transform: rotate(-1deg) !important;
            filter: brightness(0.88) contrast(1.16) saturate(1.04) drop-shadow(0 0 20px rgba(96, 165, 250, 0.34)) drop-shadow(0 0 44px rgba(56, 189, 248, 0.22)) !important;
            mix-blend-mode: normal;
            -webkit-mask-image: radial-gradient(ellipse at 72% 42%, rgba(0,0,0,1) 32%, rgba(0,0,0,0.88) 54%, rgba(0,0,0,0) 82%);
            mask-image: radial-gradient(ellipse at 72% 42%, rgba(0,0,0,1) 32%, rgba(0,0,0,0.88) 54%, rgba(0,0,0,0) 82%);
        }}

        @media (max-width: 900px) {{
            .header-container::before {{
                width: 300px !important;
                height: 300px !important;
                top: -30px !important;
                right: -52px !important;
                opacity: 0.3 !important;
            }}
        }}

        @media (max-width: 640px) {{
            .header-container::before {{
                width: 220px !important;
                height: 220px !important;
                top: -12px !important;
                right: -32px !important;
                opacity: 0.24 !important;
            }}
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )

# Initialize session state
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "knowledge_base" not in st.session_state:
    st.session_state.knowledge_base = None
if "file_name" not in st.session_state:
    st.session_state.file_name = None
if "file_type" not in st.session_state:
    st.session_state.file_type = None
if "url_summary" not in st.session_state:
    st.session_state.url_summary = None
if "image_summary" not in st.session_state:
    st.session_state.image_summary = None
if "url_source_type" not in st.session_state:
    st.session_state.url_source_type = None
if "csv_data" not in st.session_state:
    st.session_state.csv_data = None
if "csv_file" not in st.session_state:
    st.session_state.csv_file = None
if "celebrate_fx" not in st.session_state:
    st.session_state.celebrate_fx = True
if "active_sidebar_panel" not in st.session_state:
    st.session_state.active_sidebar_panel = None
if "fire_confetti" not in st.session_state:
    st.session_state.fire_confetti = False
if "welcome_confetti_done" not in st.session_state:
    st.session_state.welcome_confetti_done = False
if "download_celebration_pending" not in st.session_state:
    st.session_state.download_celebration_pending = False
if "download_celebration_nonce" not in st.session_state:
    st.session_state.download_celebration_nonce = 0
if "pending_temperature_setting" not in st.session_state:
    st.session_state.pending_temperature_setting = None
if "pending_model_setting" not in st.session_state:
    st.session_state.pending_model_setting = None
if "pending_chunk_size_setting" not in st.session_state:
    st.session_state.pending_chunk_size_setting = None
if "force_tab_index" not in st.session_state:
    st.session_state.force_tab_index = None

# Apply queued setting updates before widgets are instantiated.
if st.session_state.pending_temperature_setting is not None:
    st.session_state.temperature_setting = st.session_state.pending_temperature_setting
    st.session_state.pending_temperature_setting = None
if st.session_state.pending_model_setting is not None:
    st.session_state.model_setting = normalize_model_name(st.session_state.pending_model_setting)
    st.session_state.pending_model_setting = None
if st.session_state.pending_chunk_size_setting is not None:
    st.session_state.chunk_size_setting = st.session_state.pending_chunk_size_setting
    st.session_state.pending_chunk_size_setting = None
if "model_setting" in st.session_state:
    st.session_state.model_setting = normalize_model_name(st.session_state.model_setting)

if st.session_state.get("celebrate_fx", True) and not st.session_state.get("welcome_confetti_done", False):
    st.session_state.fire_confetti = True
    st.session_state.welcome_confetti_done = True

if st.session_state.get("celebrate_fx", True) and st.session_state.get("fire_confetti", False):
        components.html(
                """
                <script src="https://cdn.jsdelivr.net/npm/canvas-confetti@1.9.3/dist/confetti.browser.min.js"></script>
                <script>
                    const duration = 1400;
                    const end = Date.now() + duration;
                    (function frame() {
                        confetti({ particleCount: 4, angle: 60, spread: 70, origin: { x: 0 } });
                        confetti({ particleCount: 4, angle: 120, spread: 70, origin: { x: 1 } });
                        if (Date.now() < end) requestAnimationFrame(frame);
                    })();
                </script>
                """,
                height=0,
                width=0,
        )
        st.session_state.fire_confetti = False

def trigger_celebration():
    """Play celebratory effect only when enabled."""
    if st.session_state.get("celebrate_fx", True):
        st.session_state.fire_confetti = True


def trigger_download_celebration():
    """Celebrate only when the chat PDF download button is clicked."""
    if st.session_state.get("celebrate_fx", True):
        st.session_state.download_celebration_pending = True
        st.session_state.download_celebration_nonce = int(st.session_state.get("download_celebration_nonce", 0)) + 1


def reset_chat_and_source_state():
    """Clear active chat/source context and transient question inputs."""
    csv_temp_path = st.session_state.get("csv_file")
    if csv_temp_path and isinstance(csv_temp_path, str) and os.path.exists(csv_temp_path):
        try:
            os.remove(csv_temp_path)
        except OSError:
            pass

    st.session_state.chat_history = []
    st.session_state.knowledge_base = None
    st.session_state.file_name = None
    st.session_state.file_type = None
    st.session_state.url_summary = None
    st.session_state.image_summary = None
    st.session_state.url_source_type = None
    st.session_state.csv_data = None
    st.session_state.csv_file = None

    keys_to_clear = [
        "pdf_question_input", "docx_question_input", "txt_question_input",
        "json_question_input", "img_question_input", "csv_question_input", "url_question_input",
        "pdf_question_clear_pending", "docx_question_clear_pending", "txt_question_clear_pending",
        "json_question_clear_pending", "img_question_clear_pending", "csv_question_clear_pending", "url_question_clear_pending",
    ]
    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]

# Sidebar Configuration
with st.sidebar:
    st.markdown("""
    <div class="sidebar-hero" style="background: linear-gradient(135deg, rgba(49, 46, 129, 0.92) 0%, rgba(17, 24, 39, 0.96) 55%, rgba(8, 145, 178, 0.88) 100%); padding: 20px; border-radius: 18px; color: #f8fafc; text-align: center; margin-bottom: 20px; border: 1px solid rgba(255,255,255,0.08); box-shadow: 0 18px 40px rgba(2,6,23,0.35);">
        <h2 style="margin: 0; font-size: 1.5rem; color: #f8fafc;">⚙️ Configuration</h2>
        <p style="margin: 5px 0 0 0; color: #cbd5e1; font-size: 0.9rem;">Customize your AI experience</p>
    </div>
    """, unsafe_allow_html=True)

    st.toggle(
        "✨ Interactive effects",
        key="celebrate_fx",
        help="Enables celebratory effects for successful processing actions."
    )
    
    # Model Settings
    st.markdown("""
    <div class="sidebar-card" style="background: rgba(255,255,255,0.05); padding: 15px; border-radius: 14px; margin-bottom: 15px; color: #f8fafc; border: 1px solid rgba(255,255,255,0.08);">
        <h4 style="margin: 0 0 10px 0; color: #f8fafc;"><i class="fa-solid fa-microchip" style="margin-right: 7px; color: #a5b4fc;"></i>Model Settings</h4>
        <p style="margin: 0; color: #94a3b8; font-size: 0.82rem;">Tune model speed, creativity, and reasoning depth.</p>
    </div>
    """, unsafe_allow_html=True)
    if st.button("Open Model Window", use_container_width=True, key="open_model_panel"):
        st.session_state.active_sidebar_panel = "model"
    
    temperature = st.slider(
        "🌡️ Temperature (Creativity)",
        min_value=0.0,
        max_value=1.0,
        value=0.3,
        step=0.1,
        key="temperature_setting",
        help="Higher = more creative answers\nLower = more focused answers"
    )
    
    model = st.selectbox(
    "🧠 Primary Model",
    AVAILABLE_GEMINI_MODELS,
    key="model_setting"
)

    fallback_models = st.multiselect(
        "🛟 Backup Models (auto-fallback)",
        [m for m in AVAILABLE_GEMINI_MODELS if m != model],
        default=["gemini-2.0-flash"],
        key="fallback_models_setting",
        help="If primary model hits quota/rate limit, app will try these models automatically."
    )
    
    # Document Settings
    st.markdown("""
    <div class="sidebar-card" style="background: rgba(255,255,255,0.05); padding: 15px; border-radius: 14px; margin: 20px 0 15px 0; color: #f8fafc; border: 1px solid rgba(255,255,255,0.08);">
        <h4 style="margin: 0 0 10px 0; color: #f8fafc;"><i class="fa-solid fa-file-lines" style="margin-right: 7px; color: #93c5fd;"></i>Document Settings</h4>
        <p style="margin: 0; color: #94a3b8; font-size: 0.82rem;">Control chunking and retrieval sensitivity for better answers.</p>
    </div>
    """, unsafe_allow_html=True)
    if st.button("Open Document Window", use_container_width=True, key="open_doc_panel"):
        st.session_state.active_sidebar_panel = "document"
    
    chunk_size = st.slider(
        "📏 Chunk Size",
        min_value=500,
        max_value=2000,
        value=1000,
        step=100,
        key="chunk_size_setting",
        help="Size of text segments for analysis"
    )
    
    num_sources = st.slider(
        "🔍 Number of Sources",
        min_value=1,
        max_value=5,
        value=3,
        key="num_sources_setting",
        help="Relevant text chunks to retrieve"
    )
    
    # History & Data
    st.markdown("""
    <div class="sidebar-card" style="background: rgba(255,255,255,0.05); padding: 15px; border-radius: 14px; margin: 20px 0 15px 0; color: #f8fafc; border: 1px solid rgba(255,255,255,0.08);">
        <h4 style="margin: 0 0 10px 0; color: #f8fafc;"><i class="fa-solid fa-clock-rotate-left" style="margin-right: 7px; color: #93c5fd;"></i>History & Data</h4>
        <p style="margin: 0; color: #94a3b8; font-size: 0.82rem;">Manage conversations, session context, and quick stats.</p>
    </div>
    """, unsafe_allow_html=True)
    if st.button("Open History Window", use_container_width=True, key="open_history_panel"):
        st.session_state.active_sidebar_panel = "history"
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🗑️ Clear", use_container_width=True, key="clear_history"):
            reset_chat_and_source_state()
            st.rerun()
    
    with col2:
        if st.button("📊 Stats", use_container_width=True, key="show_stats"):
            with st.expander("View Statistics"):
                st.metric("Total Q&A", len(st.session_state.chat_history))
                if st.session_state.file_name:
                    st.metric("Current File", st.session_state.file_name[:30])
                if st.session_state.file_type:
                    st.metric("File Type", st.session_state.file_type.upper())
    
    st.markdown("---")
    st.markdown("""
    <div class="sidebar-meta-card" style="padding: 12px 14px; border-radius: 14px; text-align: center;">
        <span class="meta-title">SmartAI ChatBot v2.5</span>
        <span class="meta-sub">Multi-Format AI Assistant</span>
    </div>
    """, unsafe_allow_html=True)
    st.caption("📚 Powered by Gemini & LangChain")
    
    # Documentation Guide in Sidebar
    if st.button("Open Guide Window", use_container_width=True, key="open_guide_panel"):
        st.session_state.active_sidebar_panel = "guide"

    with st.sidebar.expander("📘 Application Guide", expanded=False):
        st.markdown("""
        <div class="guide-hero" style="background: linear-gradient(135deg, rgba(76,29,149,0.24), rgba(30,41,99,0.18) 42%, rgba(8,47,73,0.12)); padding: 15px; border-radius: 16px; border: 1px solid rgba(216,180,254,0.12); margin-bottom: 12px;">
            <h3 style="color: #f8fafc; margin: 0 0 8px 0;">Application Guide</h3>
            <p style="color: #cbd5e1; font-size: 13px; margin: 0; line-height: 1.7;">
                Use SmartAI ChatBot as a unified analysis workspace for documents, images, datasets, web pages, and video transcripts. Upload, index, ask, refine, and export conversations from one interface.
            </p>
            <div class="guide-chip-row">
                <span class="guide-chip">7 Input Modes</span>
                <span class="guide-chip">Live Q&A</span>
                <span class="guide-chip">PDF Export</span>
                <span class="guide-chip">URL + Video</span>
            </div>
            <div class="guide-loader">
                <div class="guide-loader-track">
                    <div class="guide-loader-node done">Source</div>
                    <div class="guide-loader-node active">Indexing</div>
                    <div class="guide-loader-node">Insights</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown('''
        <div class="guide-section-heading">
            <i class="fa-solid fa-diagram-project"></i>
            <h4>Workflow</h4>
        </div>
        ''', unsafe_allow_html=True)
        st.markdown("""
        <div class="guide-grid">
            <div class="guide-step-shell card-wine">
                <span class="guide-step-pill">Source</span>
                <strong>Choose a source format</strong>
                <div class="guide-step-body">Pick the format tab that matches your file, dataset, URL, or video. Each input path lands in the same retrieval workflow, so switching sources stays predictable.</div>
            </div>
            <div class="guide-step-shell card-teal">
                <span class="guide-step-pill">Indexing</span>
                <strong>Process and index the content</strong>
                <div class="guide-step-body">The app extracts text, splits content into chunks, and builds a searchable context layer so summaries and answers stay grounded in the uploaded source.</div>
            </div>
            <div class="guide-step-shell card-amber">
                <span class="guide-step-pill">Insights</span>
                <strong>Ask targeted questions</strong>
                <div class="guide-step-body">Use natural language to summarize, inspect, compare, and drill into facts. Hover reveals the detail so the section stays compact until you need the explanation.</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown('''
        <div class="guide-section-heading">
            <i class="fa-solid fa-layer-group"></i>
            <h4>Supported Modes</h4>
        </div>
        ''', unsafe_allow_html=True)
        tab_info = {
            "📄 PDF Explorer": "Upload PDF documents for text analysis and Q&A. Perfect for reports, articles, and ebooks.",
            "📝 Word Reader": "Process .docx files from Microsoft Word. Extract content and ask detailed questions.",
            "📄 Text Viewer": "Handle plain text files (.txt). Great for logs, notes, and simple documents.",
            "🧾 JSON Inspector": "Analyze structured JSON data. Query nested objects and arrays with ease.",
            "🖼️ Image OCR": "Extract text from images using OCR. Supports JPG, PNG, GIF, and BMP formats.",
            "📊 CSV Analyzer": "Specialized agent for spreadsheet data. Ask for summaries, statistics, and insights.",
            "🔗 URL / Video": "Process web pages and YouTube videos with transcripts, then summarize or ask follow-up questions."
        }

        mode_class_map = {
            "📄 PDF Explorer": "cap-pdf",
            "📝 Word Reader": "cap-word",
            "📄 Text Viewer": "cap-text",
            "🧾 JSON Inspector": "cap-json",
            "🖼️ Image OCR": "cap-image",
            "📊 CSV Analyzer": "cap-csv",
            "🔗 URL / Video": "cap-url",
        }
        
        for tab_name, desc in tab_info.items():
            base_class = "guide-capability featured" if "URL / Video" in tab_name or "CSV Analyzer" in tab_name else "guide-capability"
            capability_class = f"{base_class} {mode_class_map.get(tab_name, '')}".strip()
            st.markdown(f"""
            <div class="{capability_class}" style="margin: 7px 0;">
                <strong class="mode-title">{tab_name}</strong>
                <small class="mode-desc">{desc}</small>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("### ⚙️ Configuration Guide")
        
        st.markdown("""
        <div class="guide-panel">
        <h4 style="color: #f8fafc; margin-top: 0; margin-bottom: 6px;"><i class="fa-solid fa-microchip" style="margin-right: 7px; color: #a5b4fc;"></i>Model Settings</h4>
        <ul style="color: #cbd5e1; font-size: 13px; line-height: 1.7;">
            <li><strong>Temperature:</strong> Controls creativity (0.0 = focused, 1.0 = creative)</li>
            <li><strong>Model:</strong> Gemini 1.5 Flash (fast) or Gemini 1.5 Pro (deeper reasoning)</li>
        </ul>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <h4 style="color: #f8fafc; margin-bottom: 6px;"><i class="fa-solid fa-file-lines" style="margin-right: 7px; color: #93c5fd;"></i>Document Settings</h4>
        <ul style="color: #cbd5e1; font-size: 13px; line-height: 1.7;">
            <li><strong>Chunk Size:</strong> Text segment size for processing (500-2000)</li>
            <li><strong>Number of Sources:</strong> Relevant chunks to retrieve (1-5)</li>
        </ul>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <h4 style="color: #f8fafc; margin-bottom: 6px;"><i class="fa-solid fa-clock-rotate-left" style="margin-right: 7px; color: #93c5fd;"></i>History & Data</h4>
        <ul style="color: #cbd5e1; font-size: 13px; line-height: 1.7;">
            <li><strong>Clear:</strong> Reset chat history and uploaded files</li>
            <li><strong>Stats:</strong> View Q&A count and current file info</li>
        </ul>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("### 💡 Pro Tips")
        st.markdown("""
        <div class="guide-panel" style="border-color: rgba(245, 158, 11, 0.35);">
            <ul style="color: #cbd5e1; font-size: 13px; margin: 0; line-height: 1.7; padding-left: 1rem;">
                <li>📊 Use CSV Analyzer for data queries like "show me the average sales" or "plot revenue trends"</li>
                <li>🖼️ For Image OCR, ensure Tesseract is installed on your system</li>
                <li>🔍 Ask specific questions for better answers – the AI searches relevant content</li>
                <li>💾 Your chat history persists during the session but resets on page refresh</li>
                <li>⚡ Gemini Pro provides deeper responses but is typically slower than Flash</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("### 🚀 Getting Started")
        st.markdown("""
        <div class="guide-cta">
            <p style="color: #cbd5e1; font-size: 14px; margin: 0; line-height: 1.7;">
                <strong>Ready to explore?</strong> Pick a tab, upload a file, and start asking questions! 
                The AI will analyze your document and provide conversational insights.
            </p>
        </div>
        """, unsafe_allow_html=True)

# Helper functions for file processing
def extract_text_from_docx(file):
    """Extract text from DOCX file"""
    doc = Document(file)
    lines = []

    for para in doc.paragraphs:
        value = (para.text or "").strip()
        if value:
            lines.append(value)

    # Include table cells since many business DOCX files store key content in tables.
    for table in doc.tables:
        for row in table.rows:
            cell_values = []
            for cell in row.cells:
                cell_text = " ".join([(p.text or "").strip() for p in cell.paragraphs]).strip()
                if cell_text:
                    cell_values.append(cell_text)
            if cell_values:
                lines.append(" | ".join(cell_values))

    return "\n".join(lines)

def extract_text_from_txt(file):
    """Extract text from TXT file"""
    return file.read().decode("utf-8")

def extract_text_from_json(file):
    """Extract text from JSON file"""
    data = json.load(file)
    return json.dumps(data, indent=2)


def _find_tesseract_executable():
    """Return a working Tesseract executable path if available."""
    detected_path = shutil.which("tesseract")
    if detected_path:
        return detected_path

    candidates = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        os.path.join(os.environ.get("LOCALAPPDATA", ""), "Programs", "Tesseract-OCR", "tesseract.exe"),
    ]
    for candidate in candidates:
        if candidate and os.path.exists(candidate):
            return candidate
    return None


def _configure_tesseract():
    """Configure pytesseract to use a discovered Windows installation when possible."""
    tesseract_path = _find_tesseract_executable()
    if tesseract_path:
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
    return tesseract_path


def _build_tesseract_error(original_error):
    """Provide actionable OCR setup guidance."""
    details = str(original_error)
    install_steps = (
        "OCR requires Tesseract OCR to be installed on your machine. "
        "On Windows, install Tesseract and then restart Streamlit. "
        "Typical install path: C:/Program Files/Tesseract-OCR/tesseract.exe. "
        "If it is already installed, add that folder to PATH. "
        f"Original error: {details}"
    )
    return install_steps


def _extract_text_with_ocr(image):
    """Run OCR after configuring pytesseract."""
    if not _configure_tesseract():
        raise RuntimeError(_build_tesseract_error("Tesseract executable not found"))

    # Build several image variants and OCR configs to improve hit rate on scans/screenshots.
    base = image.convert("RGB")
    gray = ImageOps.grayscale(base)
    boosted = ImageEnhance.Contrast(gray).enhance(1.8)
    binary = boosted.point(lambda p: 255 if p > 160 else 0)

    variants = [base, gray, boosted, binary]
    configs = ["--oem 3 --psm 6", "--oem 3 --psm 11", "--oem 3 --psm 3"]

    best_text = ""
    try:
        for variant in variants:
            for cfg in configs:
                text = pytesseract.image_to_string(variant, config=cfg)
                if text and len(text.strip()) > len(best_text.strip()):
                    best_text = text

        if not best_text.strip() or len(best_text.strip()) < 8:
            raise RuntimeError(
                "OCR did not find enough readable text in this image. "
                "Try a clearer image, higher contrast, or crop tightly around the text."
            )

        return best_text
    except RuntimeError:
        raise
    except Exception as error:
        raise RuntimeError(_build_tesseract_error(error)) from error

def extract_text_from_image(file):
    """Extract text from image using OCR"""
    try:
        image = Image.open(file)
        text = _extract_text_with_ocr(image)
        return text
    except Exception as e:
        raise Exception(f"OCR Error: {str(e)}")


def _analyze_image_with_model(image_bytes, mime_type, primary_model, fallback_models):
    """Return a multimodal visual summary using Gemini vision-capable models."""
    if not image_bytes:
        return ""

    safe_mime = mime_type if mime_type in {"image/jpeg", "image/png", "image/gif", "image/bmp", "image/webp"} else "image/jpeg"
    image_b64 = base64.b64encode(image_bytes).decode("utf-8")

    vision_prompt = (
        "Analyze this image and provide a clear, structured explanation:\n"
        "1. Main scene/content\n"
        "2. Important objects/people and their relationships\n"
        "3. Any visible text, numbers, labels, signs, or dates\n"
        "4. Notable context (chart/diagram/screenshot/document/photo)\n"
        "5. 3 likely user questions this image can answer\n"
        "If something is unclear, state uncertainty explicitly."
    )

    def _describe_with_model(active_model):
        llm = ChatGoogleGenerativeAI(
            model=active_model,
            temperature=0.2,
            timeout=MODEL_TIMEOUT_SECONDS,
            max_retries=MODEL_MAX_RETRIES,
            google_api_key=api_key,
        )
        message = HumanMessage(
            content=[
                {"type": "text", "text": vision_prompt},
                {"type": "image_url", "image_url": f"data:{safe_mime};base64,{image_b64}"},
            ]
        )
        response = llm.invoke([message])
        return getattr(response, "content", "") or ""

    return _invoke_with_model_fallback(
        _describe_with_model,
        primary_model,
        fallback_models,
    )


def _summarize_image_content(content_text, primary_model, fallback_models):
    """Create a URL-style structured summary for image analysis output."""
    if not content_text or not content_text.strip():
        return ""

    summary_prompt = PromptTemplate(
        template="""Analyze the following image-derived content and provide:
1. A concise summary (3-5 sentences)
2. 5 key topics or themes
3. 3 suggested questions a user might ask

Content:
{text}

Response:""",
        input_variables=["text"],
    )

    def _summarize_with_model(active_model):
        llm = ChatGoogleGenerativeAI(
            temperature=0.3,
            model=active_model,
            timeout=MODEL_TIMEOUT_SECONDS,
            max_retries=MODEL_MAX_RETRIES,
            google_api_key=api_key,
        )
        chain = summary_prompt | llm | StrOutputParser()
        return chain.invoke({"text": content_text[:4000]})

    return _invoke_with_model_fallback(
        _summarize_with_model,
        primary_model,
        fallback_models,
    )

def process_file_in_chunks(uploaded_file, chunk_size_bytes=1024 * 1024):
    """Process file in chunks to avoid loading the whole file into memory."""
    file_ext = uploaded_file.name.split('.')[-1].lower()
    
    if file_ext == 'pdf':
        pdf_reader = PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            yield page.extract_text()
    elif file_ext == 'docx':
        docx_text = extract_text_from_docx(uploaded_file)
        if docx_text.strip():
            yield docx_text
    elif file_ext == 'txt':
        while True:
            raw = uploaded_file.read(chunk_size_bytes)
            if not raw:
                break
            yield raw.decode('utf-8', errors='replace')
    elif file_ext == 'json':
        data = json.load(uploaded_file)
        yield json.dumps(data, indent=2)
    elif file_ext in ['jpg', 'jpeg', 'png', 'gif', 'bmp']:
        try:
            image = Image.open(uploaded_file)
            yield _extract_text_with_ocr(image)
        except Exception as e:
            raise Exception(f"OCR Error: {str(e)}")
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")

def process_uploaded_file(uploaded_file):
    """Process uploaded file and extract text based on file type"""
    text = ""
    for chunk in process_file_in_chunks(uploaded_file):
        text += chunk
    return text


def _create_knowledge_base_from_text(raw_text, chunk_size_value):
    """Build FAISS index safely and return (index, chunk_count)."""
    if raw_text is None or not str(raw_text).strip():
        raise ValueError("No readable text was extracted from the source.")

    splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=chunk_size_value,
        chunk_overlap=200,
        length_function=len,
    )
    chunks = [chunk for chunk in splitter.split_text(str(raw_text)) if chunk and chunk.strip()]
    if not chunks:
        raise ValueError("No usable text chunks were created from the source content.")

    embeddings = get_embeddings()
    return FAISS.from_texts(chunks, embeddings), len(chunks)


def _is_gemini_quota_error(error):
    text = str(error).lower()
    quota_signals = ["resource_exhausted", "quota exceeded", "429", "generate_content_free_tier_requests"]
    return any(signal in text for signal in quota_signals)


def _is_hard_quota_exhausted(error):
    text = str(error).lower()
    hard_signals = [
        "perday",
        "per day",
        "quota exceeded for metric",
        "free_tier_requests",
    ]
    return any(signal in text for signal in hard_signals)


def _friendly_model_error_message(error):
    if _is_gemini_quota_error(error):
        return (
            "Gemini quota/rate limit reached (HTTP 429). "
            "Your free-tier limit for this model is exhausted right now. "
            "Please wait and retry, switch model/API project, or enable billing in Google AI Studio."
        )
    return f"Error: {str(error)}"


def _invoke_with_retry(invoke_fn, max_attempts=2, retry_delay_seconds=2):
    """Retry once for transient 429 errors, then raise the final exception."""
    last_error = None
    for attempt in range(max_attempts):
        try:
            return invoke_fn()
        except Exception as error:
            last_error = error
            if (
                attempt < max_attempts - 1
                and _is_gemini_quota_error(error)
                and not _is_hard_quota_exhausted(error)
            ):
                time.sleep(retry_delay_seconds)
                continue
            raise
    raise last_error


def _invoke_with_model_fallback(invoke_with_model_fn, primary_model, fallback_models):
    """Try primary model, then backups when quota errors occur."""
    last_error = None
    for model_name in _get_model_fallback_order(primary_model, fallback_models):
        try:
            return _invoke_with_retry(
                lambda: invoke_with_model_fn(model_name),
                max_attempts=1,
            )
        except Exception as error:
            last_error = error
            if _is_gemini_quota_error(error):
                continue
            raise
    raise last_error


def _build_extractive_fallback_answer(question, docs, max_sentences=4):
    """Return a lightweight extractive answer when LLM quota is exhausted."""
    joined = "\n".join([getattr(doc, "page_content", "") for doc in docs if getattr(doc, "page_content", "")])
    if not joined.strip():
        return "I could not find enough relevant text to build a fallback answer."

    sentences = re.split(r"(?<=[.!?])\s+|\n+", joined)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 30]
    if not sentences:
        return "I found relevant chunks, but not enough sentence-level text for fallback output."

    q_terms = set(re.findall(r"[a-zA-Z0-9]+", question.lower()))

    def score(sentence):
        s_terms = set(re.findall(r"[a-zA-Z0-9]+", sentence.lower()))
        overlap = len(q_terms.intersection(s_terms))
        return overlap, min(len(sentence), 220)

    ranked = sorted(sentences, key=score, reverse=True)
    top = ranked[:max_sentences]

    if not top:
        top = sentences[:max_sentences]

    bullets = "\n".join([f"- {line}" for line in top])
    return (
        "Gemini quota is exhausted right now, so this is an extractive fallback answer from your document context:\n"
        f"{bullets}"
    )

# ---------- URL / VIDEO HELPERS ----------

def _extract_youtube_id(url):
    """Extract video ID from YouTube URL variants."""
    pattern = r'(?:v=|youtu\.be\/|embed\/|shorts\/)([A-Za-z0-9_-]{11})'
    match = re.search(pattern, url)
    return match.group(1) if match else None

def _get_youtube_transcript(video_id):
    """Fetch full transcript text from a YouTube video."""
    # Support both legacy and newer youtube-transcript-api versions.
    if hasattr(YouTubeTranscriptApi, "get_transcript"):
        transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
        return " ".join([entry.get("text", "") for entry in transcript_list]).strip()

    api = YouTubeTranscriptApi()
    fetched = api.fetch(video_id)
    return " ".join([snippet.text for snippet in fetched if getattr(snippet, "text", "")]).strip()

def _scrape_web_url(url):
    """Scrape and clean text content from any web page."""
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
    response = requests.get(url, headers=headers, timeout=30)
    response.raise_for_status()
    soup = BeautifulSoup(response.content, 'html.parser')
    for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
        tag.decompose()
    lines = [line.strip() for line in soup.get_text(separator='\n').splitlines() if line.strip()]
    return '\n'.join(lines)

def render_section_heading(title, icon_class):
    st.markdown(
        f'''
        <div class="section-heading">
            <i class="fa-solid {icon_class}"></i>
            <h3>{title}</h3>
        </div>
        ''',
        unsafe_allow_html=True,
    )

# Main Header
st.markdown("""
    <div class="header-container">
        <h1>SmartAI <span class="hero-accent">ChatBot</span></h1>
        <p>Analyze documents, datasets, and multimodal sources in a unified workspace built for evidence-based research and business-grade decision support.</p>
        <div class="hero-stats">
            <div class="hero-stat">
                <strong>Multi-Format Intelligence</strong>
                <span>Ingest PDFs, DOCX, TXT, JSON, OCR images, CSV files, web URLs, and video transcripts in one streamlined workflow for cross-functional analysis.</span>
            </div>
            <div class="hero-stat">
                <strong>Conversational Retrieval</strong>
                <span>Ask iterative questions, retrieve context-aware evidence, and generate source-attributed summaries for research review and stakeholder reporting.</span>
            </div>
            <div class="hero-stat">
                <strong>Insight to Action</strong>
                <span>Transform complex source material into clear findings, practical recommendations, and shareable outputs for research and business teams.</span>
            </div>
        </div>
    </div>
""", unsafe_allow_html=True)

if st.session_state.active_sidebar_panel:
    st.markdown('<div class="hidden-close-action">', unsafe_allow_html=True)
    if st.button("Close", key="close_active_panel"):
        st.session_state.active_sidebar_panel = None
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

    if st.session_state.active_sidebar_panel == "model":
        st.markdown("""
        <div class="floating-backdrop"></div>
        <div class="floating-modal floating-modal-model">
            <div class="floating-window-bar">
                <div class="floating-window-title">Model Control Window</div>
                <div class="floating-window-actions">
                    <button class="floating-window-btn" type="button" data-win-action="close" title="Close">✕</button>
                </div>
            </div>
            <div class="floating-window-content">
                <p>Use this area to guide model behavior in real time:</p>
                <ul>
                    <li>Lower temperature for deterministic and factual responses.</li>
                    <li>Use Gemini Flash for speed and Gemini Pro for depth.</li>
                    <li>Keep prompts concrete for best retrieval quality.</li>
                </ul>
                <div class="floating-quick-actions">
                    <button class="floating-quick-btn" type="button" data-model-action="focus">Focus Mode</button>
                    <button class="floating-quick-btn" type="button" data-model-action="creative">Creative Mode</button>
                    <button class="floating-quick-btn" type="button" data-model-action="pro">Use Pro Model</button>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        st.markdown('<div class="hidden-action-bank">', unsafe_allow_html=True)
        if st.button("Focus Mode Hidden", key="focus_mode_btn"):
            st.session_state.pending_temperature_setting = 0.1
            st.rerun()
        if st.button("Creative Mode Hidden", key="creative_mode_btn"):
            st.session_state.pending_temperature_setting = 0.8
            st.rerun()
        if st.button("Use Pro Model Hidden", key="pro_model_btn"):
            st.session_state.pending_model_setting = "gemini-2.5-pro"
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
    elif st.session_state.active_sidebar_panel == "document":
        st.markdown("""
        <div class="floating-backdrop"></div>
        <div class="floating-modal">
            <div class="floating-window-bar">
                <div class="floating-window-title">Document Tuning Window</div>
                <div class="floating-window-actions">
                    <button class="floating-window-btn" type="button" data-win-action="close" title="Close">✕</button>
                </div>
            </div>
            <div class="floating-window-content">
                <p>Control chunking and context quality:</p>
                <ul>
                    <li>Smaller chunks improve pinpoint lookup.</li>
                    <li>Larger chunks preserve broader context.</li>
                    <li>Increase source count for synthesis-style answers.</li>
                </ul>
                <div class="floating-quick-actions">
                    <button class="floating-quick-btn" type="button" data-panel-action="small-chunks">Small Chunks</button>
                    <button class="floating-quick-btn" type="button" data-panel-action="balanced-chunks">Balanced</button>
                    <button class="floating-quick-btn" type="button" data-panel-action="deep-chunks">Deep Context</button>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        st.markdown('<div class="hidden-action-bank">', unsafe_allow_html=True)
        if st.button("Small Chunks Hidden", key="small_chunks_btn"):
            st.session_state.pending_chunk_size_setting = 700
            st.rerun()
        if st.button("Balanced Hidden", key="balanced_chunks_btn"):
            st.session_state.pending_chunk_size_setting = 1000
            st.rerun()
        if st.button("Deep Context Hidden", key="deep_chunks_btn"):
            st.session_state.pending_chunk_size_setting = 1600
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
    elif st.session_state.active_sidebar_panel == "history":
        st.markdown(f"""
        <div class="floating-backdrop"></div>
        <div class="floating-modal">
            <div class="floating-window-bar">
                <div class="floating-window-title">History & Session Window</div>
                <div class="floating-window-actions">
                    <button class="floating-window-btn" type="button" data-win-action="close" title="Close">✕</button>
                </div>
            </div>
            <div class="floating-window-content">
                <p>Current conversation health:</p>
                <ul>
                    <li>Total Q&A exchanges: {len(st.session_state.chat_history)}</li>
                    <li>Active source: {st.session_state.file_name if st.session_state.file_name else 'No source loaded'}</li>
                    <li>Active mode: {st.session_state.file_type if st.session_state.file_type else 'None'}</li>
                </ul>
                <div class="floating-quick-actions two-up">
                    <button class="floating-quick-btn" type="button" data-panel-action="clear-session">Clear Session</button>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        st.markdown('<div class="hidden-action-bank">', unsafe_allow_html=True)
        if st.button("Clear Session Hidden", key="clear_from_panel"):
            reset_chat_and_source_state()
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
    elif st.session_state.active_sidebar_panel == "guide":
        st.markdown("""
        <div class="floating-backdrop"></div>
        <div class="floating-modal">
            <div class="floating-window-bar">
                <div class="floating-window-title">Quick Guide Window</div>
                <div class="floating-window-actions">
                    <button class="floating-window-btn" type="button" data-win-action="close" title="Close">✕</button>
                </div>
            </div>
            <div class="floating-window-content">
                <p>Start with one of these paths:</p>
                <ul>
                    <li>Upload a file and ask direct summary and comparison questions.</li>
                    <li>Use URL / Video tab for web-page and YouTube transcript analysis.</li>
                    <li>Download final Q&A as a PDF report from the chat area.</li>
                </ul>
            </div>
        </div>
        """, unsafe_allow_html=True)

    components.html(
                """
                <script>
                (function() {
                    const doc = window.parent.document;
                    const modals = doc.querySelectorAll('.floating-modal');
                    if (!modals.length) return;

                    const hideUtilityButtons = () => {
                        const hiddenLabels = new Set([
                            'Close',
                            'Focus Mode Hidden',
                            'Creative Mode Hidden',
                            'Use Pro Model Hidden',
                            'Small Chunks Hidden',
                            'Balanced Hidden',
                            'Deep Context Hidden',
                            'Clear Session Hidden'
                        ]);

                        [...doc.querySelectorAll('button')].forEach((button) => {
                            const label = (button.textContent || '').trim();
                            if (!hiddenLabels.has(label)) return;
                            const wrapper = button.closest('[data-testid="stButton"]') || button.parentElement;
                            if (!wrapper) return;
                            wrapper.style.position = 'fixed';
                            wrapper.style.left = '-9999px';
                            wrapper.style.top = '-9999px';
                            wrapper.style.width = '1px';
                            wrapper.style.height = '1px';
                            wrapper.style.opacity = '0';
                            wrapper.style.pointerEvents = 'none';
                            wrapper.style.overflow = 'hidden';
                        });
                    };

                    hideUtilityButtons();

                    modals.forEach((modal) => {
                        if (modal.dataset.dragBound) return;
                        modal.dataset.dragBound = '1';

                        const bar = modal.querySelector('.floating-window-bar');
                        const closeBtn = modal.querySelector('[data-win-action="close"]');
                        const actionButtons = modal.querySelectorAll('[data-model-action], [data-panel-action]');
                        if (!bar || !closeBtn) return;

                        const positionInVisibleArea = () => {
                            const sidebar = doc.querySelector('[data-testid="stSidebar"]');
                            const sidebarWidth = sidebar ? Math.max(0, sidebar.getBoundingClientRect().width) : 0;
                            const vw = window.parent.innerWidth || doc.documentElement.clientWidth;
                            const modalWidth = modal.getBoundingClientRect().width;
                            const left = Math.max(10, sidebarWidth + ((vw - sidebarWidth - modalWidth) / 2));
                            modal.style.left = `${left}px`;
                            modal.style.top = '7.2rem';
                            modal.style.transform = 'none';
                        };

                        positionInVisibleArea();
                        hideUtilityButtons();

                        closeBtn.addEventListener('click', (e) => {
                            e.stopPropagation();
                            const closeActionBtn = [...doc.querySelectorAll('button')]
                                .find(b => (b.textContent || '').trim() === 'Close');
                            if (closeActionBtn) closeActionBtn.click();
                        });

                        actionButtons.forEach((actionBtn) => {
                            if (actionBtn.dataset.bound) return;
                            actionBtn.dataset.bound = '1';
                            actionBtn.addEventListener('click', (e) => {
                                e.stopPropagation();
                                const actionMap = {
                                    focus: 'Focus Mode Hidden',
                                    creative: 'Creative Mode Hidden',
                                    pro: 'Use Pro Model Hidden',
                                    'small-chunks': 'Small Chunks Hidden',
                                    'balanced-chunks': 'Balanced Hidden',
                                    'deep-chunks': 'Deep Context Hidden',
                                    'clear-session': 'Clear Session Hidden'
                                };
                                const actionKey = actionBtn.dataset.modelAction || actionBtn.dataset.panelAction;
                                const targetText = actionMap[actionKey];
                                if (!targetText) return;
                                const targetBtn = [...doc.querySelectorAll('button')]
                                    .find(b => (b.textContent || '').trim() === targetText);
                                if (targetBtn) targetBtn.click();
                            });
                        });

                        let dragging = false;
                        let offsetX = 0;
                        let offsetY = 0;

                        bar.addEventListener('mousedown', (e) => {
                            if (e.target.closest('.floating-window-btn')) return;
                            dragging = true;
                            bar.style.cursor = 'grabbing';
                            const rect = modal.getBoundingClientRect();
                            offsetX = e.clientX - rect.left;
                            offsetY = e.clientY - rect.top;
                            modal.style.transform = 'none';
                        });

                        doc.addEventListener('mousemove', (e) => {
                            if (!dragging) return;
                            const x = Math.max(8, e.clientX - offsetX);
                            const y = Math.max(8, e.clientY - offsetY);
                            modal.style.left = `${x}px`;
                            modal.style.top = `${y}px`;
                        });

                        doc.addEventListener('mouseup', () => {
                            dragging = false;
                            bar.style.cursor = 'grab';
                        });

                        window.parent.addEventListener('resize', () => {
                            if (!dragging) positionInVisibleArea();
                            hideUtilityButtons();
                        });
                    });

                })();
                </script>
                """,
                height=0,
                width=0,
    )

# Create tabs for different file types
tab_pdf, tab_docx, tab_txt, tab_json, tab_img, tab_csv, tab_url = st.tabs([
    "📄 PDF Explorer",
    "📝 Word Reader",
    "📄 Text Viewer",
    "🧾 JSON Inspector",
    "🖼️ Image OCR",
    "📊 CSV Analyzer",
    "🔗 URL / Video"
])

forced_tab_index = st.session_state.get("force_tab_index")
forced_tab_index_js = "null" if forced_tab_index is None else str(int(forced_tab_index))

# Keep selected tab stable across Streamlit reruns.
tab_persist_html = """
    <script>
    (function() {
        const doc = window.parent.document;
        const TAB_KEY = 'smartai_active_tab';
        const FORCED_TAB_INDEX = __FORCED_TAB_INDEX__;
        const EXPECTED_TABS = [
            '📄 PDF Explorer',
            '📝 Word Reader',
            '📄 Text Viewer',
            '🧾 JSON Inspector',
            '🖼️ Image OCR',
            '📊 CSV Analyzer',
            '🔗 URL / Video'
        ];
        let restoring = false;

        const getMainTabs = () => {
            const tabLists = [...doc.querySelectorAll('[data-baseweb="tab-list"]')];
            for (const tabList of tabLists) {
                const tabs = [...tabList.querySelectorAll('[data-baseweb="tab"]')];
                if (tabs.length !== EXPECTED_TABS.length) continue;
                const labels = tabs.map(t => (t.textContent || '').trim());
                const isMainMatch = EXPECTED_TABS.every(label => labels.includes(label));
                if (isMainMatch) return tabs;
            }
            return [];
        };

        const bindClicks = () => {
            getMainTabs().forEach((btn, idx) => {
                if (btn.dataset.tabPersistBound) return;
                btn.dataset.tabPersistBound = '1';
                btn.addEventListener('click', () => {
                    if (!restoring) {
                        sessionStorage.setItem(TAB_KEY, String(idx));
                    }
                });
            });
        };

        const restoreSavedTab = () => {
            const tabs = getMainTabs();
            if (!tabs.length) return;
            const savedIdx = parseInt(sessionStorage.getItem(TAB_KEY) || '0', 10);
            const targetIdx = Number.isInteger(FORCED_TAB_INDEX) ? FORCED_TAB_INDEX : savedIdx;
            if (!Number.isInteger(targetIdx) || targetIdx < 0 || targetIdx >= tabs.length) return;
            const currentIdx = tabs.findIndex(t => t.getAttribute('aria-selected') === 'true');
            if (currentIdx === targetIdx) {
                if (Number.isInteger(FORCED_TAB_INDEX)) {
                    sessionStorage.setItem(TAB_KEY, String(targetIdx));
                }
                return;
            }
            restoring = true;
            tabs[targetIdx].click();
            if (Number.isInteger(FORCED_TAB_INDEX)) {
                sessionStorage.setItem(TAB_KEY, String(targetIdx));
            }
            setTimeout(() => { restoring = false; }, 400);
        };

        const tabsNow = getMainTabs();
        if (!tabsNow.length) {
            setTimeout(() => {
                bindClicks();
                restoreSavedTab();
            }, 180);
            return;
        }

        bindClicks();
        restoreSavedTab();

        const tabList = tabsNow[0].closest('[data-baseweb="tab-list"]');
        if (!tabList) return;

        if (!tabList.dataset.persistObserverBound) {
            tabList.dataset.persistObserverBound = '1';
            new MutationObserver(() => {
                if (restoring) return;
                bindClicks();
                const tabs = getMainTabs();
                const selectedIdx = tabs.findIndex(t => t.getAttribute('aria-selected') === 'true');
                if (selectedIdx >= 0) {
                    sessionStorage.setItem(TAB_KEY, String(selectedIdx));
                }
            }).observe(tabList, { subtree: true, attributes: true, attributeFilter: ['aria-selected'] });
        }
    })();
    </script>
    """
components.html(
    tab_persist_html.replace("__FORCED_TAB_INDEX__", forced_tab_index_js),
    height=0,
    width=0,
)

# Celebrate successful PDF download clicks with a custom popper + confetti + small balloons animation.
download_confetti_nonce = int(st.session_state.get("download_celebration_nonce", 0))
download_confetti_enabled = str(
    bool(
        st.session_state.get("celebrate_fx", True)
        and st.session_state.get("download_celebration_pending", False)
    )
).lower()

download_confetti_html = """
    <script>
    (function() {
        const doc = window.parent.document;
        const STYLE_ID = 'smartai-download-celebration-style';
        const TOKEN = __DOWNLOAD_TOKEN__;
        const ENABLED = __DOWNLOAD_ENABLED__;
        const LAST_TOKEN_KEY = 'smartai_last_download_celebration';

        const ensureStyles = () => {
            if (doc.getElementById(STYLE_ID)) return;
            const style = doc.createElement('style');
            style.id = STYLE_ID;
            style.textContent = `
                @keyframes smartai-spark-pop {
                    0% {
                        transform: translate3d(0, 0, 0) scale(0.9) rotate(0deg);
                        opacity: 1;
                    }
                    100% {
                        transform: translate3d(var(--dx), var(--dy), 0) scale(0.55) rotate(var(--rot));
                        opacity: 0;
                    }
                }
                @keyframes smartai-fall {
                    0% {
                        transform: translate3d(0, -24px, 0) rotate(0deg);
                        opacity: 0;
                    }
                    12% { opacity: 1; }
                    100% {
                        transform: translate3d(var(--drift), 110vh, 0) rotate(var(--spin));
                        opacity: 0;
                    }
                }
                @keyframes smartai-balloon-rise {
                    0% {
                        transform: translate3d(0, 14px, 0) scale(0.86);
                        opacity: 0;
                    }
                    12% { opacity: 0.98; }
                    100% {
                        transform: translate3d(var(--bx), -48vh, 0) scale(1.02);
                        opacity: 0;
                    }
                }
                .smartai-celebration-layer {
                    position: fixed;
                    inset: 0;
                    pointer-events: none;
                    z-index: 99999;
                    overflow: hidden;
                }
                .smartai-spark-piece {
                    position: absolute;
                    width: 9px;
                    height: 9px;
                    border-radius: 2px;
                    box-shadow: 0 0 8px rgba(255, 255, 255, 0.55);
                    animation: smartai-spark-pop 920ms cubic-bezier(0.16, 0.84, 0.22, 1) forwards;
                }
                .smartai-fall-piece {
                    position: absolute;
                    top: -18px;
                    width: 8px;
                    height: 12px;
                    border-radius: 3px;
                    box-shadow: 0 0 10px rgba(255, 255, 255, 0.42);
                    animation: smartai-fall var(--fall) linear forwards;
                }
                .smartai-balloon {
                    position: absolute;
                    width: 24px;
                    height: 30px;
                    border-radius: 50% 50% 46% 46%;
                    animation: smartai-balloon-rise 2100ms ease-out forwards;
                    box-shadow: inset -4px -7px 10px rgba(0, 0, 0, 0.14), 0 6px 14px rgba(0, 0, 0, 0.22);
                }
                .smartai-balloon::before {
                    content: '';
                    position: absolute;
                    left: 6px;
                    top: 6px;
                    width: 6px;
                    height: 9px;
                    border-radius: 999px;
                    background: rgba(255, 255, 255, 0.65);
                    transform: rotate(-25deg);
                }
                .smartai-balloon::after {
                    content: '';
                    position: absolute;
                    left: 11px;
                    bottom: -14px;
                    width: 1.6px;
                    height: 16px;
                    background: rgba(255, 255, 255, 0.72);
                }
            `;
            doc.head.appendChild(style);
        };

        const spawnBurst = (layer, originX, originY, count) => {
            const colors = ['#2dd4bf', '#22d3ee', '#facc15', '#f43f5e', '#60a5fa', '#a78bfa'];

            for (let i = 0; i < count; i += 1) {
                const p = doc.createElement('span');
                p.className = 'smartai-spark-piece';
                p.style.left = `${originX}px`;
                p.style.top = `${originY}px`;
                p.style.background = colors[Math.floor(Math.random() * colors.length)];

                const angle = Math.random() * Math.PI * 2;
                const distance = 70 + Math.random() * 190;
                const dx = Math.cos(angle) * distance;
                const dy = Math.sin(angle) * distance + 70;
                const rot = `${Math.floor(Math.random() * 720 - 360)}deg`;

                p.style.setProperty('--dx', `${dx}px`);
                p.style.setProperty('--dy', `${dy}px`);
                p.style.setProperty('--rot', rot);
                p.style.opacity = '0.98';
                p.style.width = `${6 + Math.random() * 7}px`;
                p.style.height = `${5 + Math.random() * 7}px`;
                p.style.animationDelay = `${Math.random() * 120}ms`;

                layer.appendChild(p);
            }
        };

        const spawnShinyRain = (layer) => {
            const colors = ['#2dd4bf', '#facc15', '#f472b6', '#60a5fa', '#34d399'];
            for (let i = 0; i < 70; i += 1) {
                const p = doc.createElement('span');
                p.className = 'smartai-fall-piece';
                p.style.left = `${Math.random() * 100}vw`;
                p.style.background = colors[Math.floor(Math.random() * colors.length)];
                p.style.setProperty('--drift', `${Math.floor(Math.random() * 180 - 90)}px`);
                p.style.setProperty('--spin', `${Math.floor(Math.random() * 900 - 450)}deg`);
                p.style.setProperty('--fall', `${1400 + Math.random() * 1500}ms`);
                p.style.animationDelay = `${Math.random() * 260}ms`;
                layer.appendChild(p);
            }
        };

        const spawnSmallBalloons = (layer) => {
            const colors = ['#22d3ee', '#f472b6', '#fbbf24'];
            const count = 2 + Math.floor(Math.random() * 2);
            const cx = Math.max(160, Math.min(window.parent.innerWidth - 160, window.parent.innerWidth * 0.5));
            for (let i = 0; i < count; i += 1) {
                const b = doc.createElement('span');
                b.className = 'smartai-balloon';
                b.style.left = `${cx + (i - 1) * 44 + (Math.random() * 22 - 11)}px`;
                b.style.top = `${window.parent.innerHeight - 110 + Math.random() * 12}px`;
                b.style.background = colors[i % colors.length];
                b.style.setProperty('--bx', `${Math.floor(Math.random() * 90 - 45)}px`);
                b.style.animationDelay = `${80 + i * 110}ms`;
                layer.appendChild(b);
            }
        };

        const fireCelebration = () => {
            const layer = doc.createElement('div');
            layer.className = 'smartai-celebration-layer';
            const cx = Math.round(window.parent.innerWidth * 0.5);
            const cy = Math.round(window.parent.innerHeight * 0.30);

            spawnBurst(layer, cx - 120, cy, 42);
            spawnBurst(layer, cx + 120, cy, 42);
            setTimeout(() => spawnBurst(layer, cx, cy - 18, 52), 110);
            spawnShinyRain(layer);
            spawnSmallBalloons(layer);

            doc.body.appendChild(layer);
            setTimeout(() => layer.remove(), 3400);
        };

        ensureStyles();

        if (!ENABLED || TOKEN <= 0) return;
        const lastToken = parseInt(sessionStorage.getItem(LAST_TOKEN_KEY) || '0', 10);
        if (Number.isInteger(lastToken) && lastToken >= TOKEN) return;

        sessionStorage.setItem(LAST_TOKEN_KEY, String(TOKEN));
        fireCelebration();
    })();
    </script>
"""
components.html(
    download_confetti_html
        .replace("__DOWNLOAD_TOKEN__", str(download_confetti_nonce))
        .replace("__DOWNLOAD_ENABLED__", download_confetti_enabled),
    height=0,
    width=0,
)

if st.session_state.get("download_celebration_pending", False):
    st.session_state.download_celebration_pending = False

if st.session_state.get("force_tab_index") is not None:
    st.session_state.force_tab_index = None

# helper to render file info card

def _display_file_info(uploaded, extra_desc=""):
    size_bytes = uploaded.size
    if size_bytes >= 1024 ** 3:
        size_str = f"{size_bytes / 1024 ** 3:.2f} GB"
    elif size_bytes >= 1024 ** 2:
        size_str = f"{size_bytes / 1024 ** 2:.1f} MB"
    else:
        size_str = f"{size_bytes / 1024:.1f} KB"
    render_section_heading("File Info", "fa-circle-info")
    safe_name = _html.escape(uploaded.name[:25])
    extra_html = extra_desc if extra_desc else ""
    html_block = (
        f'<div class="metric-card">'
        f'<strong>📄 Name:</strong> {safe_name}...<br>'
        f'<strong>📊 Size:</strong> {size_str}<br>'
        f'{extra_html}'
        f'<strong>🔖 Status:</strong> ✓ Ready'
        f'</div>'
    )
    st.markdown(html_block, unsafe_allow_html=True)


def _extract_agent_output(result):
    """Normalize CSV agent output into a readable string."""
    if isinstance(result, dict):
        for key in ("output", "answer", "final_answer", "result"):
            value = result.get(key)
            if value is not None and str(value).strip():
                return str(value)
        return str(result)
    return str(result)


def _csv_fallback_answer(user_question, csv_path, llm):
    """Fallback answer when the CSV agent fails or returns empty output."""
    df = pd.read_csv(csv_path)
    rows, cols = df.shape
    schema_text = "\n".join([f"- {col}: {dtype}" for col, dtype in df.dtypes.items()])
    sample_csv = df.head(20).to_csv(index=False)

    numeric_cols = df.select_dtypes(include=["number"]).columns.tolist()
    stats_text = "No numeric columns detected."
    if numeric_cols:
        stats_text = df[numeric_cols].describe().to_string()

    fallback_template = """You are a data analyst helping with CSV questions.

Dataset info:
- Rows: {rows}
- Columns: {cols}

Schema:
{schema}

Sample rows (first 20):
{sample}

Numeric summary:
{stats}

User question:
{question}

Answer clearly. If exact computation needs all rows beyond the sample, state that and provide the best possible estimate using the available context.
"""
    prompt = PromptTemplate(
        template=fallback_template,
        input_variables=["rows", "cols", "schema", "sample", "stats", "question"],
    )
    chain = prompt | llm | StrOutputParser()
    return chain.invoke(
        {
            "rows": rows,
            "cols": cols,
            "schema": schema_text,
            "sample": sample_csv,
            "stats": stats_text,
            "question": user_question,
        }
    )

# generic chat UI for a given file_type filter
def _chat_interface(filter_type, question_key, question_label):
    render_section_heading("Chat Interface", "fa-comments")
    input_key = f"{question_key}_input"
    clear_input_key = f"{question_key}_clear_pending"
    if input_key not in st.session_state:
        st.session_state[input_key] = ""
    if clear_input_key not in st.session_state:
        st.session_state[clear_input_key] = False
    if st.session_state[clear_input_key]:
        st.session_state[input_key] = ""
        st.session_state[clear_input_key] = False

    def _set_suggested_question(suggested_text):
        st.session_state[input_key] = suggested_text

    should_show = False
    if filter_type == "csv":
        should_show = st.session_state.get("csv_data") is not None
    else:
        should_show = (
            st.session_state.knowledge_base is not None
            and st.session_state.get("file_type") == filter_type
        )

    messages_for_type = [msg for msg in st.session_state.chat_history if msg.get("file_type") == filter_type]

    SUGGESTED = {
        "pdf":   ["Summarize this document", "What are the key points?", "List the main topics covered", "What conclusions does this document make?"],
        "docx":  ["Summarize this document", "What are the key points?", "List the main topics covered", "What are the action items?"],
        "txt":   ["Summarize this text", "What are the key points?", "What topics are discussed?", "Give a brief overview"],
        "json":  ["What fields are in this data?", "Describe the structure of this JSON", "What are the top-level keys?", "Summarize the data"],
        "image": ["What text is in this image?", "Describe the content", "List any numbers or dates found", "Summarize the image text"],
        "csv":   ["Show me the first few rows", "What columns are available?", "Give a statistical summary", "What trends do you see?"],
        "url":   ["Summarize this page", "What are the key points?", "List the main topics", "What is this page about?"],
    }

    if messages_for_type:
        rendered_messages = []
        for message in messages_for_type:
            safe_question = _html.escape(str(message.get("question", "")))
            safe_answer = _html.escape(str(message.get("answer", ""))).replace("\n", "<br>")
            rendered_messages.append(f'<div class="message-user"><strong>You:</strong> {safe_question}</div>')
            rendered_messages.append(f'<div class="message-assistant"><strong>Bot:</strong> {safe_answer}</div>')

        st.markdown(
            f'<div class="chat-container">{"".join(rendered_messages)}</div>',
            unsafe_allow_html=True,
        )
    elif should_show:
        suggestions = SUGGESTED.get(filter_type, [])
        st.markdown("""
        <p style="margin: 0.5rem 0 0.8rem 0; font-size:0.85rem; color:var(--text-2); font-weight:600; letter-spacing:0.04em;">
            ✨ SUGGESTED QUESTIONS
        </p>""", unsafe_allow_html=True)
        for i, suggestion in enumerate(suggestions):
            st.button(
                suggestion,
                key=f"suggest_{filter_type}_{i}",
                use_container_width=True,
                on_click=_set_suggested_question,
                args=(suggestion,),
            )

    st.markdown("---")
    
    with st.form(key=f"{filter_type}_chat_form", clear_on_submit=True):
        user_question = st.text_input(
            question_label,
            placeholder="Type your question here...",
            key=input_key,
            disabled=not should_show
        )
        send_clicked = st.form_submit_button(
            "Send Message",
            use_container_width=True,
            disabled=not should_show,
        )

    if send_clicked:
        if user_question and user_question.strip():
            if filter_type == "csv":
                with st.spinner("📊 Analyzing..."):
                    try:
                        if not HAS_TABULATE:
                            st.error("Missing optional dependency 'tabulate'. Install it using: pip install tabulate")
                            return
                        if not api_key:
                            st.error("Missing GOOGLE_API_KEY. Add it to your .env file to use CSV analysis.")
                            return
                        def _ask_csv_with_model(active_model):
                            llm = ChatGoogleGenerativeAI(
                                model=active_model,
                                temperature=temperature,
                                google_api_key=api_key,
                                timeout=MODEL_TIMEOUT_SECONDS,
                                max_retries=MODEL_MAX_RETRIES,
                            )
                            try:
                                agent = create_csv_agent(
                                    llm,
                                    st.session_state.csv_file,
                                    verbose=False,
                                    allow_dangerous_code=True,
                                )
                                result = agent.invoke({"input": user_question})
                                response_text = _extract_agent_output(result)
                                if response_text and response_text.strip():
                                    return response_text
                            except Exception as agent_error:
                                if not _is_gemini_quota_error(agent_error):
                                    st.warning(
                                        f"CSV agent issue detected on {active_model}, using fallback analyzer. "
                                        f"Details: {str(agent_error)}"
                                    )
                                else:
                                    raise

                            return _csv_fallback_answer(user_question, st.session_state.csv_file, llm)

                        response = _invoke_with_model_fallback(
                            _ask_csv_with_model,
                            model,
                            fallback_models,
                        )
                    except Exception as e:
                        st.error(_friendly_model_error_message(e))
                        return
            else:
                with st.spinner("🤔 Thinking..."):
                    docs = []
                    try:
                        docs = st.session_state.knowledge_base.similarity_search(user_question, k=num_sources)
                        if not docs:
                            st.warning("❌ No relevant information found.")
                            return
                        template = """Use the following context to answer the question. If you don't know, say so.

{context}

Question: {question}
Answer:"""
                        context = "\n\n---\n\n".join([doc.page_content for doc in docs])
                        prompt = PromptTemplate(template=template, input_variables=["context", "question"])

                        def _ask_docs_with_model(active_model):
                            llm = ChatGoogleGenerativeAI(
                                temperature=temperature,
                                model=active_model,
                                timeout=MODEL_TIMEOUT_SECONDS,
                                max_retries=MODEL_MAX_RETRIES,
                            )
                            chain = prompt | llm | StrOutputParser()
                            return chain.invoke({"context": context, "question": user_question})

                        response = _invoke_with_model_fallback(
                            _ask_docs_with_model,
                            model,
                            fallback_models,
                        )
                    except Exception as e:
                        if _is_gemini_quota_error(e) and docs:
                            st.info("Gemini quota exceeded. Using local extractive fallback answer.")
                            response = _build_extractive_fallback_answer(user_question, docs)
                        else:
                            st.error(_friendly_model_error_message(e))
                            return
            entry = {"timestamp": datetime.now().isoformat(),
                     "question": user_question,
                     "answer": response,
                     "file_type": filter_type}
            st.session_state.chat_history.append(entry)
            st.session_state[clear_input_key] = True
            tab_index_map = {
                "pdf": 0,
                "docx": 1,
                "txt": 2,
                "json": 3,
                "image": 4,
                "csv": 5,
                "url": 6,
            }
            st.session_state.force_tab_index = tab_index_map.get(filter_type)
            st.rerun()
        elif user_question:
            st.warning("⚠️ Please enter a valid question.")

    st.markdown("---")
    has_messages = len(messages_for_type) > 0
    pdf_bytes = b""
    if has_messages:
        summary_for_export = None
        if filter_type == "url":
            summary_for_export = st.session_state.get("url_summary")
        elif filter_type == "image":
            summary_for_export = st.session_state.get("image_summary")
        pdf_bytes = generate_chat_pdf(
            messages_for_type,
            file_name=st.session_state.get("file_name"),
            summary_text=summary_for_export,
        )
    fname = f"SmartAI_Chat_{filter_type}_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
    st.download_button(
        label="Download Chat PDF",
        data=pdf_bytes,
        file_name=fname,
        mime="application/pdf",
        use_container_width=True,
        disabled=not has_messages,
        key=f"download_pdf_{filter_type}",
        on_click=trigger_download_celebration,
    )
    if not has_messages:
        st.caption("Add at least one Q&A in this tab to enable PDF download.")

# ---------- PDF TAB ----------
with tab_pdf:
    col_upload, col_chat = st.columns([1, 2])
    with col_upload:
        render_section_heading("Upload PDF", "fa-file-arrow-up")
        uploaded_pdf = st.file_uploader("Choose a PDF", type=["pdf"], key="pdf_uploader")
        if uploaded_pdf is not None:
            process_pdf_clicked = st.button("🔄 Process PDF", use_container_width=True, key="process_pdf_btn")
            if process_pdf_clicked:
                if st.session_state.file_name == uploaded_pdf.name and st.session_state.file_type == "pdf" and st.session_state.knowledge_base is not None:
                    st.info("This PDF is already processed and ready for questions.")
                else:
                    with st.spinner("🔄 Processing PDF..."):
                        try:
                            text = process_uploaded_file(uploaded_pdf)
                            st.session_state.knowledge_base, _ = _create_knowledge_base_from_text(text, chunk_size)
                            st.session_state.file_name = uploaded_pdf.name
                            st.session_state.file_type = "pdf"
                            _display_file_info(uploaded_pdf)
                            trigger_celebration()
                        except Exception as e:
                            st.error(f"❌ Error: {str(e)}")
    with col_chat:
        _chat_interface("pdf", "pdf_question", "Ask about the PDF:")

# ---------- WORD TAB ----------
with tab_docx:
    col_upload, col_chat = st.columns([1, 2])
    with col_upload:
        render_section_heading("Upload Word Document", "fa-file-word")
        uploaded_doc = st.file_uploader("Choose a DOCX file", type=["docx"], key="docx_uploader")
        if uploaded_doc is not None:
            process_docx_clicked = st.button("🔄 Process DOCX", use_container_width=True, key="process_docx_btn")
            if process_docx_clicked:
                if st.session_state.file_name == uploaded_doc.name and st.session_state.file_type == "docx" and st.session_state.knowledge_base is not None:
                    st.info("This DOCX is already processed and ready for questions.")
                else:
                    with st.spinner("🔄 Processing document..."):
                        try:
                            text = process_uploaded_file(uploaded_doc)
                            st.session_state.knowledge_base, _ = _create_knowledge_base_from_text(text, chunk_size)
                            st.session_state.file_name = uploaded_doc.name
                            st.session_state.file_type = "docx"
                            _display_file_info(uploaded_doc)
                            trigger_celebration()
                        except Exception as e:
                            st.error(f"❌ Error: {str(e)}")
    with col_chat:
        _chat_interface("docx", "docx_question", "Ask about the Word doc:")

# ---------- TEXT TAB ----------
with tab_txt:
    col_upload, col_chat = st.columns([1, 2])
    with col_upload:
        render_section_heading("Upload Text File", "fa-file-lines")
        uploaded_txt = st.file_uploader("Choose a TXT file", type=["txt"], key="txt_uploader")
        if uploaded_txt is not None:
            process_txt_clicked = st.button("🔄 Process TXT", use_container_width=True, key="process_txt_btn")
            if process_txt_clicked:
                if st.session_state.file_name == uploaded_txt.name and st.session_state.file_type == "txt" and st.session_state.knowledge_base is not None:
                    st.info("This TXT file is already processed and ready for questions.")
                else:
                    with st.spinner("🔄 Processing text..."):
                        try:
                            text = process_uploaded_file(uploaded_txt)
                            st.session_state.knowledge_base, _ = _create_knowledge_base_from_text(text, chunk_size)
                            st.session_state.file_name = uploaded_txt.name
                            st.session_state.file_type = "txt"
                            _display_file_info(uploaded_txt)
                            trigger_celebration()
                        except Exception as e:
                            st.error(f"❌ Error: {str(e)}")
    with col_chat:
        _chat_interface("txt", "txt_question", "Ask about the text:")

# ---------- JSON TAB ----------
with tab_json:
    col_upload, col_chat = st.columns([1, 2])
    with col_upload:
        render_section_heading("Upload JSON File", "fa-code")
        uploaded_json = st.file_uploader("Choose a JSON file", type=["json"], key="json_uploader")
        if uploaded_json is not None:
            process_json_clicked = st.button("🔄 Process JSON", use_container_width=True, key="process_json_btn")
            if process_json_clicked:
                if st.session_state.file_name == uploaded_json.name and st.session_state.file_type == "json" and st.session_state.knowledge_base is not None:
                    st.info("This JSON file is already processed and ready for questions.")
                else:
                    with st.spinner("🔄 Processing JSON..."):
                        try:
                            text = process_uploaded_file(uploaded_json)
                            st.session_state.knowledge_base, _ = _create_knowledge_base_from_text(text, chunk_size)
                            st.session_state.file_name = uploaded_json.name
                            st.session_state.file_type = "json"
                            _display_file_info(uploaded_json)
                            trigger_celebration()
                        except Exception as e:
                            st.error(f"❌ Error: {str(e)}")
    with col_chat:
        _chat_interface("json", "json_question", "Ask about the JSON data:")

# ---------- IMAGE TAB ----------
with tab_img:
    col_upload, col_chat = st.columns([1, 2])
    with col_upload:
        render_section_heading("Upload Image", "fa-image")
        if not _find_tesseract_executable():
            st.warning(
                "Tesseract OCR is not installed. Install it on Windows and restart the app. "
                "Expected path is usually C:/Program Files/Tesseract-OCR/tesseract.exe"
            )
        uploaded_img = st.file_uploader("Choose an image", type=["jpg","jpeg","png","gif","bmp"], key="img_uploader")
        if uploaded_img is not None:
            process_image_clicked = st.button("🔄 Process Image", use_container_width=True, key="process_image_btn")
            if process_image_clicked:
                if st.session_state.file_name == uploaded_img.name and st.session_state.file_type == "image" and st.session_state.knowledge_base is not None:
                    st.info("This image is already processed and ready for questions.")
                else:
                    with st.spinner("🔄 Processing image..."):
                        try:
                            image_bytes = uploaded_img.getvalue()
                            ocr_text = extract_text_from_image(io.BytesIO(image_bytes))
                            visual_summary = ""

                            if api_key:
                                with st.spinner("🧠 Running visual analysis with Gemini..."):
                                    try:
                                        visual_summary = _analyze_image_with_model(
                                            image_bytes,
                                            uploaded_img.type,
                                            model,
                                            fallback_models,
                                        )
                                    except Exception as vision_error:
                                        st.warning(
                                            "Visual analysis fallback engaged. OCR is still available. "
                                            f"Details: {_friendly_model_error_message(vision_error)}"
                                        )
                            else:
                                st.info("GOOGLE_API_KEY not found. Running OCR-only mode for image analysis.")

                            merged_text_parts = [f"OCR Extracted Text:\n{ocr_text.strip()}"]
                            if visual_summary and visual_summary.strip():
                                merged_text_parts.append(f"Visual Analysis Summary:\n{visual_summary.strip()}")
                            merged_text = "\n\n".join(merged_text_parts)

                            image_summary = visual_summary
                            if api_key:
                                try:
                                    image_summary = _summarize_image_content(
                                        merged_text,
                                        model,
                                        fallback_models,
                                    )
                                except Exception as summary_error:
                                    st.warning(
                                        "Image summary generation fallback engaged. "
                                        f"Details: {_friendly_model_error_message(summary_error)}"
                                    )

                            st.session_state.knowledge_base, _ = _create_knowledge_base_from_text(merged_text, chunk_size)
                            st.session_state.file_name = uploaded_img.name
                            st.session_state.file_type = "image"
                            st.session_state.image_summary = image_summary or None
                            _display_file_info(uploaded_img)
                            trigger_celebration()
                        except Exception as e:
                            st.error(f"❌ Error: {str(e)}")

        if st.session_state.get("file_type") == "image" and st.session_state.get("image_summary"):
            render_section_heading("Image AI Summary", "fa-wand-magic-sparkles")
            st.markdown(
                f"""
                <div style="background: rgba(255,255,255,0.04); padding: 15px; border-radius: 14px; border-left: 4px solid #22d3ee; font-size: 14px; color: #e2e8f0; border: 1px solid rgba(255,255,255,0.08); line-height: 1.8;">
                {st.session_state.image_summary.replace(chr(10), '<br>')}
                </div>
                """,
                unsafe_allow_html=True,
            )
    with col_chat:
        _chat_interface("image", "img_question", "Ask about the image:")

# ---------- CSV TAB ----------
with tab_csv:
    col_upload, col_chat = st.columns([1, 2])
    with col_upload:
        render_section_heading("Upload CSV File", "fa-table")
        uploaded_csv = st.file_uploader("Select a CSV file", type="csv", key="csv_uploader")
        if uploaded_csv is not None:
            process_csv_clicked = st.button("🔄 Process CSV", use_container_width=True, key="process_csv_btn")
            if process_csv_clicked:
                if st.session_state.file_name == uploaded_csv.name and st.session_state.file_type == "csv":
                    st.info("This CSV is already processed and ready for questions.")
                else:
                    st.session_state.file_name = uploaded_csv.name
                    st.session_state.file_type = "csv"
                    with st.spinner("🔄 Processing CSV..."):
                        try:
                            # Save the uploaded file to a portable temporary file
                            tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.csv')
                            tmp.write(uploaded_csv.getbuffer())
                            tmp.flush()
                            temp_file_path = tmp.name
                            tmp.close()

                            # Store the path in session state
                            st.session_state.csv_file = temp_file_path

                            # Read the first few rows to display info
                            df = pd.read_csv(temp_file_path, nrows=5)
                            st.session_state.csv_data = df

                            _display_file_info(uploaded_csv, f"<br><strong>📊 Dimensions:</strong> Previewing first 5 rows<br>")
                            trigger_celebration()
                        except Exception as e:
                            st.error(f"Error: {str(e)}")

    with col_chat:
        _chat_interface("csv", "csv_question", "Ask about your data:")

# ---------- URL / VIDEO TAB ----------
with tab_url:
    col_input, col_chat = st.columns([1, 2])
    with col_input:
        render_section_heading("Enter URL", "fa-link")
        st.markdown("""
        <div style="background: rgba(255,255,255,0.04); padding: 12px; border-radius: 14px; border-left: 4px solid #06b6d4; margin-bottom: 15px; border: 1px solid rgba(255,255,255,0.08);">
            <strong style="color: #67e8f9;">Supported Sources</strong><br>
            <small style="color: #cbd5e1;">
                📺 YouTube videos (with captions/subtitles)<br>
                🌐 Web articles, blogs &amp; news pages<br>
                📰 Documentation &amp; online reports
            </small>
        </div>
        """, unsafe_allow_html=True)

        url_input = st.text_input(
            "Paste URL here",
            placeholder="https://www.youtube.com/watch?v=... or any web page",
            key="url_input"
        )

        if st.button("🔄 Process URL", use_container_width=True, key="process_url_btn"):
            if url_input and url_input.strip():
                with st.spinner("🔍 Fetching & analyzing content..."):
                    try:
                        video_id = _extract_youtube_id(url_input)
                        if video_id:
                            source_type = "YouTube Video"
                            raw_text = _get_youtube_transcript(video_id)
                        else:
                            source_type = "Web Page"
                            raw_text = _scrape_web_url(url_input)

                        if not raw_text or len(raw_text.strip()) < 100:
                            st.error("❌ Could not extract meaningful content from this URL.")
                        else:
                            st.session_state.knowledge_base, chunk_count = _create_knowledge_base_from_text(raw_text, chunk_size)
                            st.session_state.file_name = url_input
                            st.session_state.file_type = "url"
                            st.session_state.url_source_type = source_type

                            # Auto-generate NotebookLM-style summary
                            with st.spinner("📝 Generating summary & key topics..."):
                                summary_prompt = PromptTemplate(
                                    template="""Analyze the following content and provide:
1. A concise summary (3-5 sentences)
2. 5 key topics or themes
3. 3 suggested questions a user might ask

Content:
{text}

Response:""",
                                    input_variables=["text"]
                                )

                                def _summarize_with_model(active_model):
                                    llm = ChatGoogleGenerativeAI(
                                        temperature=0.3,
                                        model=active_model,
                                        timeout=MODEL_TIMEOUT_SECONDS,
                                        max_retries=MODEL_MAX_RETRIES,
                                    )
                                    chain = summary_prompt | llm | StrOutputParser()
                                    return chain.invoke({"text": raw_text[:4000]})

                                st.session_state.url_summary = _invoke_with_model_fallback(
                                    _summarize_with_model,
                                    model,
                                    fallback_models,
                                )

                            st.success(f"✅ {source_type} processed! {chunk_count} chunks indexed.")
                            trigger_celebration()
                            st.session_state.force_tab_index = 6
                            st.rerun()
                    except Exception as e:
                        st.error(f"❌ {_friendly_model_error_message(e)}")
            else:
                st.warning("⚠️ Please enter a URL first.")

        if st.session_state.get("url_summary"):
            render_section_heading("AI Summary", "fa-wand-magic-sparkles")
            st.markdown(f"""
            <div style="background: rgba(255,255,255,0.04); padding: 15px; border-radius: 14px; border-left: 4px solid #10b981; font-size: 14px; color: #e2e8f0; border: 1px solid rgba(255,255,255,0.08); line-height: 1.8;">
            {st.session_state.url_summary.replace(chr(10), '<br>')}
            </div>
            """, unsafe_allow_html=True)

    with col_chat:
        _chat_interface("url", "url_question", "Ask about this content:")

# Footer
st.markdown("---")
col1, col2, col3 = st.columns(3)
with col1:
    st.caption("📚 Powered by Gemini")
with col2:
    st.caption("🔐 Secure & Private")
with col3:
    st.caption("⚡ Streamlit & LangChain")
